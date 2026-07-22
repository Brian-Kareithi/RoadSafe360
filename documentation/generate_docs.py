#!/usr/bin/env python3
"""Generate comprehensive technical documentation for RoadSafe360 as .docx."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Global styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    h.font.name = 'Calibri'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# ── Helper functions ──
def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light List Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    doc.add_paragraph()

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(0.5)
    return p

def bold(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_toc():
    doc.add_paragraph()
    p = doc.add_paragraph('(Table of Contents will be generated when you open in Word and update fields)')
    p.italic = True
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RoadSafe360')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Intelligent Driver Demerit & Road Safety Management System')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('Complete Technical Documentation & Handover Guide')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'Document Generated: {datetime.datetime.now().strftime("%B %d, %Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info2.add_run('For: PBL Classroom Project | Kenya Road Safety Context')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
add_toc()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)

doc.add_heading('1.1 Purpose of the Application', level=2)
doc.add_paragraph(
    'RoadSafe360 is a full-stack web application for "Intelligent Driver Demerit & Road Safety Management." '
    'It simulates how countries like Kenya manage driver licences using demerit points, suspensions, enforcement '
    'actions, and analytics. The application is built as a PBL (Project-Based Learning) classroom project and '
    'provides four distinct user roles: Admin, Police Officer, Driver, and Road Authority.'
)
doc.add_paragraph(
    'Core capabilities include: digital driver licence issuance with QR code verification, traffic offence recording '
    'with GPS location and evidence upload, demerit points tracking and automatic suspension enforcement, '
    'appeal submission and review workflow, role-specific dashboards with KPIs and analytics, report generation '
    '(PDF/CSV), regional road safety statistics, system settings management, and notification delivery.'
)

doc.add_heading('1.2 High-Level Architecture', level=2)
doc.add_paragraph(
    'The application follows a client-rendered single-page application (SPA) architecture using Next.js App Router '
    'with React 19. All pages are client-side rendered (\'use client\'). Firebase serves as the backend platform, '
    'providing Authentication (email/password), Firestore (NoSQL database), Storage (file uploads for evidence/photos), '
    'and Cloud Functions (for server-side logic).'
)
doc.add_paragraph(
    'Data flows directly between the React frontend and Firebase services via the Firebase Web SDK. There are no '
    'custom API route handlers or backend servers — the Firestore security rules enforce row-level access control. '
    'The application uses React Context for state management (authentication and theme), a custom hook for Firestore '
    'real-time subscriptions (useFirestore), and service modules that wrap Firebase SDK calls.'
)

doc.add_heading('1.3 Technology Stack', level=2)
add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['Next.js', '16.2.10', 'React framework with App Router for page routing and SSR/SSG'],
        ['React', '19.2.4', 'UI component library for building the user interface'],
        ['TypeScript', '5.x', 'Type-safe JavaScript for improved developer experience'],
        ['Tailwind CSS', '4.x', 'Utility-first CSS framework for rapid UI development'],
        ['Firebase Auth', '12.15.0', 'Authentication service (email/password)'],
        ['Firebase Firestore', '12.15.0', 'NoSQL document database for all application data'],
        ['Firebase Storage', '12.15.0', 'File storage for evidence photos and driver images'],
        ['Firebase Functions', '12.15.0', 'Serverless functions (configured but not actively used)'],
        ['Recharts', '3.9.2', 'Charting library for dashboard analytics and visualizations'],
        ['Leaflet / react-leaflet', '1.9.4 / 5.0.0', 'OpenStreetMap integration for GPS location picking'],
        ['jsPDF / html2canvas', '4.2.1 / 1.4.1', 'PDF generation for reports and digital licence export'],
        ['react-qr-code', '2.2.0', 'QR code generation for licence verification'],
        ['react-hot-toast', '2.6.0', 'Toast notification library for user feedback'],
        ['Radix UI', 'Various', 'Headless UI primitives (Avatar, Dialog, Select, etc.)'],
        ['date-fns', '4.4.0', 'Date formatting and manipulation'],
        ['react-hook-form', '7.81.0', 'Form state management (installed but sparingly used)'],
        ['Zod', '4.4.3', 'Schema validation (installed but sparingly used)'],
        ['lucide-react / react-icons', 'Various', 'Icon libraries for UI elements'],
    ]
)

doc.add_heading('1.4 Major Dependencies and Their Roles', level=2)
deps = [
    ('firebase', 'Provides the entire backend: authentication, NoSQL database, file storage, and serverless functions. All persistent data flows through Firestore.'),
    ('next', 'Provides the React framework with file-based routing via the App Router. Handles compilation, bundling, and development server.'),
    ('tailwindcss', 'Provides utility-first CSS classes. All styling in the application uses Tailwind classes directly in JSX. No separate CSS files except globals.css.'),
    ('recharts', 'Provides the chart components (BarChart, PieChart, LineChart, AreaChart) used across admin dashboard, authority dashboard, driver dashboard, and analytics page.'),
    ('leaflet / react-leaflet', 'Provides interactive map functionality for the offence-issuance workflow. Used in MapPicker component to let officers select GPS location of an offence.'),
    ('jspdf / html2canvas', 'Enables PDF generation from HTML elements. Used in the digital licence page (download as PDF) and reports page (generate PDF reports).'),
    ('react-qr-code', 'Generates QR codes for the digital driver licence. Encodes licence number, national ID, and status for verification.'),
    ('@radix-ui/*', 'Provides accessible, unstyled UI primitives used as foundations for the shadcn-style component library.'),
    ('class-variance-authority', 'Enables the variant-based component styling pattern used in Button, Badge, and other UI primitives.'),
    ('clsx / tailwind-merge', 'Utility libraries for conditional class name merging. The cn() function combines them for clean class composition.'),
]
for dep, purpose in deps:
    p = doc.add_paragraph()
    run = p.add_run(f'{dep}: ')
    run.bold = True
    p.add_run(purpose)

doc.add_heading('1.5 Folder Structure', level=2)
folders = [
    ('app/', 'Next.js App Router pages. Contains all route segments as directories with page.tsx files. Each subdirectory represents a route path.'),
    ('components/', 'Reusable React components. Contains the dashboard shell (Sidebar, TopBar), UI primitives (Button, Card, Badge, etc.), and specialized components (MapPicker, AuthGuard, FirstTimeTutorial, PWAInstallPrompt, SkeletonLoader, ServiceWorkerRegister). Organised into ui/ for atomic design-system components.'),
    ('components/ui/', 'Shadcn-style UI primitives. Atomic components: avatar, badge, button, card, input, select, table. Use class-variance-authority for variant management.'),
    ('contexts/', 'React Context providers for global state: AuthContext (authentication state, user profile, role, logout) and ThemeContext (light/dark theme toggle with localStorage persistence).'),
    ('hooks/', 'Custom React hooks. Contains useFirestore.ts with useCollection, useDocument, addDocument, updateDocument, deleteDocument, and getDocument for Firestore CRUD operations with real-time subscriptions.'),
    ('firebase/', 'Firebase SDK initialization and configuration. Exports auth, db, storage, and functions instances.'),
    ('lib/', 'Utility libraries: format.ts (currency, date, status color, risk level, demerit status formatting) and utils.ts (cn() function for class merging).'),
    ('services/', 'Firebase service wrapper modules: authService.ts (login, register, logout, profile fetching), driverService.ts (search, getByLicence), and offenceService.ts (CRUD operations for offences and categories).'),
    ('types/', 'TypeScript type definitions. Single index.ts file containing all interfaces: Driver, Licence, PoliceOfficer, OffenceCategory, OffenceRecord, Appeal, Suspension, AuditLog, Notification, SystemSetting, Vehicle, Payment, QRCodeData, AppUser, UserRole.'),
    ('seed/', 'Database seed scripts executed with tsx. seed.ts creates all Firebase Auth users, Firestore collections, and sample data. add-offences.ts adds matatu/nganya-specific offence categories.'),
    ('public/', 'Static assets served at the root URL. Contains logo.png, favicon set in favicon_io/, SVG icons, PWA manifest.json, and site.webmanifest.'),
    ('utils/', 'Contains a single cn.ts file — duplicates lib/utils.ts with identical functionality. This is dead code.'),
    ('documentation/', 'Documentation output directory. Contains this file and the generation script.'),
]
add_table(['Directory', 'Purpose'], folders)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. APPLICATION FLOW
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Application Flow', level=1)

doc.add_heading('2.1 Application Startup', level=2)
doc.add_paragraph(
    'The application is a Next.js App Router project. On startup, the root layout (app/layout.tsx) renders. '
    'It imports the Inter font from Google Fonts, sets metadata (title, description, icons, PWA manifest), '
    'and wraps all children in <ClientLayout> which provides ThemeProvider, AuthProvider, DashboardShell (for '
    'authenticated pages), PWAInstallPrompt, and a Toaster. The root page (app/page.tsx) immediately redirects '
    'to /auth via router.replace().'
)

doc.add_heading('2.2 Initialization Process', level=2)
doc.add_paragraph(
    'When a user visits any page: (1) ThemeProvider checks localStorage for "rs-theme" key and applies the dark '
    'class to <html>. (2) AuthProvider subscribes to Firebase Auth state changes via onAuthStateChanged. '
    '(3) When auth state resolves, the user\'s profile is fetched from Firestore (users/{uid}). (4) AuthGuard '
    'checks authentication — if not authenticated, redirects to /auth. (5) If authenticated, the role-specific '
    'dashboard renders with the sidebar navigation filtered by user role.'
)

doc.add_heading('2.3 Routing', level=2)
doc.add_paragraph(
    'Routing uses Next.js App Router conventions. All pages are marked \'use client\' for client-side rendering. '
    'The route structure is flat: /auth, /dashboard/{role}, /offences, /offences/new, /appeals, /drivers, '
    '/licence, /reports, /analytics, /regions, /notifications, /settings. The root page (/) redirects to /auth. '
    'The dashboard layout (app/dashboard/layout.tsx) is a pass-through wrapper. ClientLayout provides the '
    'DashboardShell wrapper for all routes except /auth and /.'
)

doc.add_heading('2.4 Authentication Flow', level=2)
doc.add_paragraph(
    'Authentication uses Firebase Auth with email/password. The flow: (1) User enters credentials on /auth page '
    'or clicks a quick-login button. (2) handleLogin or quickLogin calls login() from authService, which wraps '
    'signInWithEmailAndPassword. (3) On success, the AuthProvider\'s onAuthStateChanged listener fires, setting '
    'the user state. (4) getUserProfile fetches the user document from Firestore (users/{uid}) to get role and '
    'displayName. (5) The useEffect in AuthPage detects the authenticated user and redirects to the role-specific '
    'dashboard. (6) AuthGuard on protected pages ensures unauthenticated users are redirected to /auth.'
)

doc.add_heading('2.5 Authorization', level=2)
doc.add_paragraph(
    'Authorization is role-based with four roles: admin, police, driver, authority. The role is stored in the '
    'Firestore users document and loaded into the AuthContext profile. Authorization checks occur at two levels: '
    '(1) The Sidebar component filters navigation links based on the user\'s role — each role sees only their '
    'permitted pages. (2) Firestore security rules (firestore.rules) enforce document-level access control on the '
    'database side, using helper functions isAdmin(), isPolice(), isDriver(), isAuthority(), and isOwner().'
)

doc.add_heading('2.6 State Management', level=2)
doc.add_paragraph(
    'State management uses React built-in features: (1) React Context for global state — AuthContext holds '
    'user, profile, loading, role, and logout. ThemeContext holds theme, toggle, and setTheme. (2) Custom hooks '
    'for data fetching — useCollection and useDocument from useFirestore.ts use Firestore\'s onSnapshot for '
    'real-time subscriptions. (3) Local component state with useState for form inputs, UI state, and ephemeral '
    'data. (4) No external state management library (Redux, Zustand, etc.) is used.'
)

doc.add_heading('2.7 Data Flow: Frontend to Backend', level=2)
doc.add_paragraph(
    'Data flows from frontend to Firebase and back as follows:'
)
doc.add_paragraph(
    'WRITE FLOW: User interacts with a form (e.g., issuing an offence) → component calls an async function from '
    'useFirestore (addDocument) or a service module → the function calls Firebase SDK (addDoc) → Firestore security '
    'rules validate the request → data is written to Firestore → the operation completes → the component handles '
    'the response (toast notification, form reset, state update).'
)
doc.add_paragraph(
    'READ FLOW: Component mounts → useCollection or useDocument hook initializes → onSnapshot establishes a '
    'real-time listener on a Firestore query → Firestore evaluates security rules → initial data snapshot arrives → '
    'hook sets data state → component re-renders with data → any subsequent data changes trigger automatic '
    're-renders via the snapshot listener.'
)
doc.add_paragraph(
    'AUTH FLOW: User logs in → Firebase Auth SDK authenticates → onAuthStateChanged fires → AuthProvider fetches '
    'user profile from Firestore → profile data is available via useAuth() hook throughout the component tree.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. FOLDER-BY-FOLDER BREAKDOWN
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Folder-by-Folder Breakdown', level=1)

folder_details = [
    {
        'name': 'app/',
        'purpose': 'Next.js App Router pages creating the application\'s route structure.',
        'responsibilities': 'Contains all page components, one per route. Each subdirectory represents a URL segment and contains a page.tsx file.',
        'dependencies': 'All components, hooks, contexts, services, types, lib utilities.',
        'files': 'layout.tsx, page.tsx, auth/page.tsx, dashboard/layout.tsx, dashboard/admin/page.tsx, dashboard/authority/page.tsx, dashboard/driver/page.tsx, dashboard/police/page.tsx, offences/page.tsx, offences/new/page.tsx, appeals/page.tsx, drivers/page.tsx, licence/page.tsx, reports/page.tsx, analytics/page.tsx, regions/page.tsx, notifications/page.tsx, settings/page.tsx, globals.css, favicon.ico.',
        'interactions': 'Root layout wraps all pages in ClientLayout. Each page uses useCollection/useDocument hooks to read Firestore data and renders UI components from components/. Pages call service functions directly for auth mutations.',
    },
    {
        'name': 'components/',
        'purpose': 'Reusable React components forming the UI library and application shell.',
        'responsibilities': 'Provides the dashboard layout shell (Sidebar + TopBar wrapper), UI primitives (Button, Card, Badge, Table, Input, Select, Avatar), authentication guard, onboarding tutorial, map picker, skeleton loader, PWA install prompt, and service worker registration.',
        'dependencies': 'contexts/AuthContext, contexts/ThemeContext, lib/utils, react-icons, leaflet, react-qr-code, jspdf, html2canvas, react-hot-toast.',
        'files': 'AuthGuard.tsx, ClientLayout.tsx, DashboardShell.tsx, FirstTimeTutorial.tsx, MapPicker.tsx, PWAInstallPrompt.tsx, ServiceWorkerRegister.tsx, Sidebar.tsx, SkeletonLoader.tsx, TopBar.tsx, ui/avatar.tsx, ui/badge.tsx, ui/button.tsx, ui/card.tsx, ui/input.tsx, ui/select.tsx, ui/table.tsx.',
        'interactions': 'DashboardShell wraps all authenticated pages. Sidebar uses AuthContext for role-based navigation. TopBar uses ThemeContext for theme toggle. MapPicker uses Leaflet for GPS location in offence issuance. AuthGuard guards protected routes.',
    },
    {
        'name': 'contexts/',
        'purpose': 'React Context providers for global application state.',
        'responsibilities': 'AuthContext manages Firebase auth state, user profile, role, and logout. ThemeContext manages light/dark theme with localStorage persistence.',
        'dependencies': 'firebase/config (for AuthContext), services/authService, types/index.',
        'files': 'AuthContext.tsx, ThemeContext.tsx.',
        'interactions': 'AuthProvider wraps the entire app in layout.tsx. ThemeProvider wraps the entire app. Both contexts are consumed by Sidebar, TopBar, ClientLayout, and page components throughout the app.',
    },
    {
        'name': 'hooks/',
        'purpose': 'Custom React hooks for data fetching and Firestore operations.',
        'responsibilities': 'Provides real-time Firestore collection and document subscriptions via onSnapshot, plus async CRUD operations (add, update, delete, get).',
        'dependencies': 'firebase/config (db), firebase/firestore SDK.',
        'files': 'useFirestore.ts.',
        'interactions': 'Used by every page component that reads or writes data. useCollection provides real-time data with loading and error states. addDocument/updateDocument/deleteDocument are used in service modules and directly in page components.',
    },
    {
        'name': 'lib/',
        'purpose': 'Utility and formatting functions.',
        'responsibilities': 'cn() function for Tailwind class merging. Formatting functions for currency (KES), dates (en-KE locale), status colors, risk labels, and demerit point status.',
        'dependencies': 'clsx, tailwind-merge.',
        'files': 'utils.ts, format.ts.',
        'interactions': 'cn() is used throughout all components. Format functions are used in every data-displaying page (dashboards, tables, cards).',
    },
    {
        'name': 'services/',
        'purpose': 'Firebase SDK wrapper modules providing a clean API for auth and data operations.',
        'responsibilities': 'authService wraps Firebase Auth operations (sign in, sign up, sign out, profile fetch). driverService wraps driver search queries. offenceService wraps offence CRUD and category fetching.',
        'dependencies': 'firebase/config, types/index, firebase/auth, firebase/firestore.',
        'files': 'authService.ts, driverService.ts, offenceService.ts.',
        'interactions': 'authService is consumed by AuthContext. driverService and offenceService are consumed by page components (driver lookup, offence issuance, offence listing).',
    },
    {
        'name': 'types/',
        'purpose': 'TypeScript type definitions for all data models.',
        'responsibilities': 'Defines interfaces for Driver, Licence, PoliceOfficer, OffenceCategory, OffenceRecord, Appeal, Suspension, AuditLog, Notification, SystemSetting, Vehicle, Payment, QRCodeData, AppUser, and UserRole type.',
        'dependencies': 'None (pure TypeScript types).',
        'files': 'index.ts.',
        'interactions': 'Imported by services, hooks, contexts, and page components to provide type safety for Firestore data operations.',
    },
    {
        'name': 'firebase/',
        'purpose': 'Firebase SDK initialization and configuration.',
        'responsibilities': 'Initializes Firebase app with hardcoded configuration values. Exports auth, db, storage, and functions instances.',
        'dependencies': 'firebase/app, firebase/auth, firebase/firestore, firebase/storage, firebase/functions.',
        'files': 'config.ts.',
        'interactions': 'Imported by hooks/useFirestore, services/*, and seed scripts. All Firestore and Auth operations use these instances.',
    },
    {
        'name': 'seed/',
        'purpose': 'Database population scripts for development and demo.',
        'responsibilities': 'seed.ts creates 14 Auth users, 14 driver profiles with licences, 2 police officers, 14 offence categories, 20+ sample offences, appeals, 8 regions, 10 system settings, and notifications. add-offences.ts adds 2 matatu/nganya-specific offence categories.',
        'dependencies': 'firebase/app, firebase/firestore, firebase/auth.',
        'files': 'seed.ts, add-offences.ts.',
        'interactions': 'Standalone scripts executed via npm run seed and npm run seed:offences. Not imported by the application.',
    },
    {
        'name': 'public/',
        'purpose': 'Static assets served at the root URL path.',
        'responsibilities': 'Hosts the application logo (logo.png), favicon set for various devices and PWA, SVG icons (file, globe, next, vercel, window), PWA manifest (manifest.json), and site.webmanifest.',
        'dependencies': 'None (static files).',
        'files': 'logo.png, manifest.json, site.webmanifest, favicon_io/*, file.svg, globe.svg, next.svg, vercel.svg, window.svg.',
        'interactions': 'logo.png is referenced in layout.tsx metadata, auth page, sidebar, and skeleton loader. manifest.json is linked in layout.tsx for PWA support. Favicons are configured in layout.tsx metadata.',
    },
    {
        'name': 'utils/',
        'purpose': 'Contains a duplicate of the cn() utility function.',
        'responsibilities': 'Duplicates lib/utils.ts with identical code. This appears to be dead code — a leftover from refactoring.',
        'dependencies': 'clsx, tailwind-merge.',
        'files': 'cn.ts.',
        'interactions': 'None — the app imports cn() from lib/utils.ts. This file is unused.',
    },
]

for fd in folder_details:
    doc.add_heading(f'3.{folder_details.index(fd)+1} {fd["name"]}', level=2)
    bold(fd['purpose'])
    doc.add_paragraph(f'Responsibilities: {fd["responsibilities"]}')
    doc.add_paragraph(f'Dependencies: {fd["dependencies"]}')
    doc.add_paragraph(f'Files: {fd["files"]}')
    doc.add_paragraph(f'Interactions: {fd["interactions"]}')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. FILE DOCUMENTATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. File Documentation', level=1)

files_doc = [
    # App files
    ('app/layout.tsx', 'Root layout component. Sets HTML structure, imports Inter font, configures PWA metadata, and wraps content in ClientLayout.', [
        ('Exports', 'RootLayout (default)'),
        ('Metadata', 'title, description, icons, appleWebApp, manifest'),
        ('Dependencies', 'next/font/google, @/components/ClientLayout, globals.css'),
        ('Usage', 'Entry point for the application. Every page renders through this layout.'),
    ]),
    ('app/page.tsx', 'Home page that immediately redirects to /auth.', [
        ('Exports', 'Home (default)'),
        ('Logic', 'useEffect triggers router.replace(\'/auth\') on mount. Returns null (no UI).'),
        ('Dependencies', 'next/navigation (useRouter), react (useEffect)'),
        ('Called by', 'Next.js router when user visits /'),
    ]),
    ('app/globals.css', 'Global CSS with Tailwind CSS v4 directives, CSS custom properties for theming, animations, and Leaflet styles.', [
        ('CSS Variables', '--primary, --secondary, --accent, --bg, --text, --sidebar-*, --shadow-*, --radius-*. Dark mode overrides in .dark class.'),
        ('Animations', 'fadeIn, fadeInUp, scaleIn, slideInRight, pulse-dot, shimmer, countUp'),
        ('Stagger Classes', 'stagger-1 through stagger-10 (animation-delay increments of 40ms)'),
        ('Utilities', 'Custom scrollbar, Leaflet container styling'),
    ]),
    ('app/auth/page.tsx', 'Login page with email/password form and quick-login buttons for demo accounts.', [
        ('Exports', 'AuthPage (default)'),
        ('State', 'email, password, showPw, submitting'),
        ('Key Functions', 'handleLogin (form submit), quickLogin (demo button click)'),
        ('Behavior', 'Redirects authenticated users to role-specific dashboard via useEffect. Shows SkeletonLoader during loading.'),
        ('Dependencies', 'AuthContext, authService, UI components (Button, Input, Card), react-hot-toast, react-icons'),
    ]),
    ('app/dashboard/layout.tsx', 'Pass-through layout that renders children with no modifications.', [
        ('Exports', 'DashboardLayout (default)'),
        ('Usage', 'Simply wraps dashboard page content. The actual layout is handled by DashboardShell in ClientLayout.'),
    ]),
    ('app/dashboard/admin/page.tsx', 'Admin dashboard showing national overview stats, monthly offences bar chart, severity pie chart, and recent activity.', [
        ('Exports', 'AdminDashboard (default)'),
        ('Data Sources', 'useCollection(\'drivers\'), useCollection(\'licences\'), useCollection(\'offences\'), useCollection(\'policeOfficers\'), useCollection(\'appeals\')'),
        ('Key Calculations', 'suspended/revoked/high-risk counts, today offences, total revenue, monthly trends, severity distribution'),
        ('Dependencies', 'recharts (BarChart, PieChart), format.ts, UI components (Card, Badge), react-icons'),
    ]),
    ('app/dashboard/authority/page.tsx', 'Authority dashboard with national KPIs, offence trends charts, revenue trends, and risk prediction overview.', [
        ('Exports', 'AuthorityDashboard (default)'),
        ('Data Sources', 'useCollection(\'drivers\'), useCollection(\'licences\'), useCollection(\'offences\'), useCollection(\'appeals\')'),
        ('KPIs', 'Suspension rate, appeal approval rate, avg points per offence, high-risk drivers'),
        ('Dependencies', 'recharts (AreaChart, BarChart), format.ts, UI components (Card)'),
    ]),
    ('app/dashboard/driver/page.tsx', 'Driver dashboard showing demerit points gauge, offence history, appeals count, notifications, and digital licence link.', [
        ('Exports', 'DriverDashboard (default)'),
        ('Data Sources', 'useAuth, useCollection(\'offences\'), useCollection(\'appeals\'), useCollection(\'notifications\'), useCollection(\'drivers\')'),
        ('Key Features', 'Pie chart showing remaining vs used points, demerit status badge, recent offences list, summary cards'),
        ('Dependencies', 'recharts (PieChart), format.ts, AuthContext, UI components (Card, Badge, Button)'),
    ]),
    ('app/dashboard/police/page.tsx', 'Police dashboard showing today\'s cases, driver search input, stats cards, and recent offences list.', [
        ('Exports', 'PoliceDashboard (default)'),
        ('Data Sources', 'useCollection(\'offences\'), useCollection(\'appeals\')'),
        ('Key Features', 'Global search bar for driver lookup, today cases count, quick-issue offence button'),
        ('Dependencies', 'UI components (Card, Badge, Button, Input), format.ts'),
    ]),
    ('app/offences/page.tsx', 'Offences list page with search, status filter, and sortable table.', [
        ('Exports', 'OffencesPage (default)'),
        ('State', 'search (string), statusFilter (string)'),
        ('Logic', 'Client-side filtering by search term and status. Sort by timestamp descending.'),
        ('Dependencies', 'useCollection(\'offences\'), Table, Badge, Input, Select, format.ts'),
    ]),
    ('app/offences/new/page.tsx', 'Create new offence page with driver lookup, category selector, GPS map picker, notes, and evidence upload buttons.', [
        ('Exports', 'NewOffencePage (default)'),
        ('State', 'driverId, categoryId, gpsLocation, notes, submitting, driverInfo, loadingDriver, driverNotFound'),
        ('Key Functions', 'lookupDriver (debounced Firestore lookup), handleDriverChange, handleSubmit'),
        ('Features', 'Auto-verify driver on input blur (500ms debounce), selected category preview, lazy-loaded MapPicker, photo/evidence upload buttons (UI only)'),
        ('Dependencies', 'useCollection(\'offenceCategories\'), MapPicker (lazy), useFirestore (dynamic import), toast, UI components'),
    ]),
    ('app/appeals/page.tsx', 'Appeals list and submission page with inline form.', [
        ('Exports', 'AppealsPage (default)'),
        ('State', 'showForm, reason, offenceId'),
        ('Key Functions', 'handleSubmit (creates appeal document)'),
        ('Dependencies', 'useCollection(\'appeals\'), useCollection(\'offences\'), useFirestore (dynamic import), UI components'),
    ]),
    ('app/drivers/page.tsx', 'Driver management page with searchable, sortable table.', [
        ('Exports', 'DriversPage (default)'),
        ('State', 'search (string)'),
        ('Logic', 'Client-side filter by name, national ID, or phone. Points badge color varies by balance.'),
        ('Dependencies', 'useCollection(\'drivers\'), Table, Badge, Input, format.ts (risk label/color)'),
    ]),
    ('app/licence/page.tsx', 'Digital driver licence card with QR code, vehicle listings, PDF download, and print support.', [
        ('Exports', 'DigitalLicencePage (default)'),
        ('Data Sources', 'useAuth, useCollection(\'drivers\'), useCollection(\'licences\'), useCollection(\'vehicles\')'),
        ('Key Functions', 'downloadPDF (html2canvas + jsPDF), printLicence (window.open print)'),
        ('Features', 'QR code generation (react-qr-code), registered vehicles list, status-aware card styling, printable HTML template'),
        ('Dependencies', 'react-qr-code, jspdf, html2canvas, toast'),
    ]),
    ('app/reports/page.tsx', 'Report generation page supporting 11 report types with PDF and CSV export.', [
        ('Exports', 'ReportsPage (default)'),
        ('State', 'selectedReport'),
        ('Key Functions', 'generatePDF (jsPDF with header/body/footer), handleExportCSV, getReportData'),
        ('Report Types', 'daily, weekly, monthly, annual, regional, driver-history, officer-activity, revenue, suspension, high-risk, repeat-offenders'),
        ('Dependencies', 'useCollection (drivers, offences, officers), format.ts, jspdf'),
    ]),
    ('app/analytics/page.tsx', 'Analytics page with interactive charts showing offence trends, revenue, licence status distribution, and risk distribution.', [
        ('Exports', 'AnalyticsPage (default)'),
        ('Data Sources', 'useCollection(\'offences\'), useCollection(\'drivers\'), useCollection(\'licences\')'),
        ('Charts', 'LineChart (monthly offence trends), BarChart (revenue trends), PieChart (licence status), BarChart (risk distribution)'),
        ('Dependencies', 'recharts (LineChart, BarChart, PieChart), format.ts'),
    ]),
    ('app/regions/page.tsx', 'Regional road safety statistics page with hardcoded region data and interactive map placeholder.', [
        ('Exports', 'RegionsPage (default)'),
        ('Data', 'Static array of 8 Kenyan regions with driver counts, offence counts, stations, and trends'),
        ('Features', 'Region cards with stats, interactive map placeholder'),
        ('Dependencies', 'UI components (Card, Badge), react-icons'),
    ]),
    ('app/notifications/page.tsx', 'Notification centre with type-based icons, time stamps, read/unread styling, and mark-all-read functionality.', [
        ('Exports', 'NotificationsPage (default)'),
        ('State', 'notifications (from useCollection)'),
        ('Key Functions', 'markAllRead (iterates and updates each notification)'),
        ('Dependencies', 'useCollection(\'notifications\'), useFirestore (dynamic import), format.ts, UI components (Card, Badge, Button)'),
    ]),
    ('app/settings/page.tsx', 'System settings page for demerit thresholds and general configuration.', [
        ('Exports', 'SettingsPage (default)'),
        ('State', 'formValues (Record<string, string>)'),
        ('Key Functions', 'handleSave (iterates settings, updates changed values)'),
        ('Dependencies', 'useCollection(\'settings\'), useFirestore (dynamic import), UI components (Card, Input, Button, Select)'),
    ]),
    # Components
    ('components/AuthGuard.tsx', 'Route guard component that redirects unauthenticated users to /auth.', [
        ('Exports', 'AuthGuard'),
        ('Props', '{ children: React.ReactNode }'),
        ('Logic', 'If loading, shows SkeletonLoader. If no user/profile, redirects to /auth. Otherwise renders children.'),
        ('Dependencies', 'AuthContext, SkeletonLoader, next/navigation'),
        ('Used by', 'DashboardShell'),
    ]),
    ('components/ClientLayout.tsx', 'Client-side layout wrapper providing theme, auth, dashboard shell, PWA prompt, and toaster.', [
        ('Exports', 'ClientLayout'),
        ('Logic', 'Determines if current route is /auth or / to decide whether to wrap in DashboardShell. Unregisters old service workers on mount.'),
        ('Dependencies', 'ThemeContext, AuthContext, DashboardShell, PWAInstallPrompt, react-hot-toast'),
        ('Used by', 'app/layout.tsx'),
    ]),
    ('components/DashboardShell.tsx', 'Dashboard layout shell that wraps authenticated page content with Sidebar, TopBar, and AuthGuard.', [
        ('Exports', 'DashboardShell'),
        ('Props', '{ children: React.ReactNode }'),
        ('State', 'mobileMenuOpen'),
        ('Dependencies', 'Sidebar, TopBar, FirstTimeTutorial, AuthGuard'),
    ]),
    ('components/Sidebar.tsx', 'Role-based navigation sidebar with links, user info, and logout button.', [
        ('Exports', 'Sidebar'),
        ('Props', '{ open: boolean, onClose: () => void }'),
        ('Data', 'navLinks record mapping roles to link arrays. Uses AuthContext for role, profile, logout.'),
        ('Logic', 'Highlights active link based on pathname. Closes on navigation in mobile view. Shows user avatar and name at bottom.'),
        ('Dependencies', 'AuthContext, lib/utils (cn), next/navigation, next/link, react-icons'),
    ]),
    ('components/TopBar.tsx', 'Top header bar with page title, notification bell, theme toggle, and user avatar.', [
        ('Exports', 'TopBar'),
        ('Props', '{ onMenuToggle: () => void, mobileMenuOpen: boolean }'),
        ('Logic', 'Extracts page name from pathname. Shows theme icon based on current theme.'),
        ('Dependencies', 'ThemeContext, AuthContext, next/navigation, react-icons'),
    ]),
    ('components/MapPicker.tsx', 'Leaflet-based interactive map for selecting GPS coordinates.', [
        ('Exports', 'MapPicker (default)'),
        ('Props', '{ value: string, onChange: (value: string) => void }'),
        ('State', 'mapReady, address, locating'),
        ('Key Functions', 'init (async Leaflet initialization), updateValue (reverse geocoding via Nominatim), locateMe (browser geolocation)'),
        ('Dependencies', 'leaflet, leaflet/dist/leaflet.css (dynamic imports), react-icons'),
        ('Used by', 'app/offences/new/page.tsx (lazy loaded)'),
    ]),
    ('components/SkeletonLoader.tsx', 'Full-page skeleton loading screen with application logo, shimmer effects, and placeholder cards.', [
        ('Exports', 'SkeletonLoader'),
        ('Features', 'Logo display, animated shimmer bars, placeholder card grid, border-accented sections'),
        ('Dependencies', 'globals.css animations (animate-shimmer)'),
    ]),
    ('components/FirstTimeTutorial.tsx', 'Onboarding tutorial overlay shown on first visit per role.', [
        ('Exports', 'FirstTimeTutorial'),
        ('State', 'step, visible'),
        ('Key Functions', 'dismiss (stores in localStorage), step navigation'),
        ('Logic', 'Checks localStorage for rs-tutorial-{role}. Shows 3 steps per role with skip/done options.'),
        ('Dependencies', 'AuthContext, lib/utils (cn), react-icons'),
    ]),
    ('components/PWAInstallPrompt.tsx', 'Progressive Web App install prompt that captures the beforeinstallprompt event.', [
        ('Exports', 'PWAInstallPrompt'),
        ('State', 'deferredPrompt, showPrompt, isStandalone'),
        ('Key Functions', 'handleInstall (shows native install prompt)'),
        ('Logic', 'Hides if already running in standalone mode (PWA). Shows install button with app logo.'),
    ]),
    ('components/ServiceWorkerRegister.tsx', 'Service worker registration component.', [
        ('Exports', 'ServiceWorkerRegister'),
        ('Logic', 'Attempts to register /sw.js service worker on mount. Silently catches errors.'),
        ('Note', 'No /sw.js file exists in the public directory. This component is not imported anywhere.'),
    ]),
    # Contexts
    ('contexts/AuthContext.tsx', 'Authentication context providing user, profile, loading, role, and logout throughout the component tree.', [
        ('Exports', 'AuthProvider, useAuth'),
        ('State', 'user (Firebase User | null), profile (AppUser | null), loading (boolean), role (UserRole | null)'),
        ('Logic', 'Subscribes to onAuthStateChanged. On auth change, fetches user profile from Firestore. Logout function calls authService.logout and redirects to /auth.'),
        ('Dependencies', 'services/authService, firebase/auth (User type), types/index (AppUser, UserRole)'),
        ('Consumed by', 'AuthGuard, Sidebar, TopBar, FirstTimeTutorial, driver dashboard, licence page, auth page'),
    ]),
    ('contexts/ThemeContext.tsx', 'Theme context providing light/dark mode with localStorage persistence.', [
        ('Exports', 'ThemeProvider, useTheme'),
        ('State', 'theme (\'light\' | \'dark\'), mounted'),
        ('Logic', 'On mount, reads localStorage rs-theme key or prefers-color-scheme media query. Toggles .dark class on document.documentElement.'),
        ('Consumed by', 'TopBar (toggle button), ClientLayout (wrapper)'),
    ]),
    # Hooks
    ('hooks/useFirestore.ts', 'Custom hooks and functions for Firestore real-time reads and CRUD operations.', [
        ('Exports', 'useCollection, useDocument, addDocument, updateDocument, deleteDocument, getDocument'),
        ('useCollection', 'Real-time subscription to a Firestore collection with optional query constraints. Returns { data, loading, error }.'),
        ('useDocument', 'Real-time subscription to a single document. Returns { data, loading }.'),
        ('addDocument', 'Async function that adds a document with createdAt/updatedAt timestamps.'),
        ('updateDocument', 'Async function that updates a document with updatedAt timestamp.'),
        ('deleteDocument', 'Async function that deletes a document.'),
        ('getDocument', 'Async function for one-time document fetch (not real-time).'),
        ('Dependencies', 'firebase/firestore (collection, query, onSnapshot, doc, getDoc, addDoc, updateDoc, deleteDoc), firebase/config (db)'),
    ]),
    # Services
    ('services/authService.ts', 'Authentication service wrapping Firebase Auth operations.', [
        ('Exports', 'onAuthChange, login, register, logout, getUserProfile'),
        ('onAuthChange', 'Subscribes to Firebase auth state changes. Returns unsubscribe function.'),
        ('login', 'Signs in with email and password via signInWithEmailAndPassword.'),
        ('register', 'Creates user account and writes user profile to Firestore.'),
        ('logout', 'Signs out via signOut.'),
        ('getUserProfile', 'Fetches user document from Firestore users/{uid} collection.'),
        ('Dependencies', 'firebase/config (auth, db), firebase/auth, firebase/firestore, types/index'),
    ]),
    ('services/driverService.ts', 'Driver search and lookup service.', [
        ('Exports', 'searchDrivers, getDriverByLicence'),
        ('searchDrivers', 'Queries drivers collection by name prefix. Returns matching Driver array.'),
        ('getDriverByLicence', 'Queries drivers collection by licenceNumber. Returns single Driver or null.'),
        ('Dependencies', 'firebase/config (db), firebase/firestore, types/index'),
    ]),
    ('services/offenceService.ts', 'Offence management service.', [
        ('Exports', 'createOffence, getOffencesByDriver, getOffenceCategories'),
        ('createOffence', 'Creates a new offence document with timestamp.'),
        ('getOffencesByDriver', 'Queries offences by driverId, ordered by timestamp descending.'),
        ('getOffenceCategories', 'Fetches all offence category documents.'),
        ('Dependencies', 'firebase/config (db), firebase/firestore, types/index'),
    ]),
    # Lib
    ('lib/format.ts', 'Formatting utility functions for display.', [
        ('Exports', 'formatCurrency, formatDate, formatDateTime, getStatusColor, getRiskColor, getRiskLabel, getDemeritStatus'),
        ('formatCurrency', 'Formats number as KES currency using Intl.NumberFormat with en-KE locale.'),
        ('formatDate', 'Formats date as "Jan 1, 2024" using en-KE locale.'),
        ('formatDateTime', 'Formats date as "Jan 1, 2024, 02:30 PM" using en-KE locale.'),
        ('getStatusColor', 'Returns Tailwind CSS classes for status badges based on status string.'),
        ('getRiskColor', 'Returns text color class based on risk score (0-1).'),
        ('getRiskLabel', 'Returns text label for risk score (Low/Medium/High/Unknown).'),
        ('getDemeritStatus', 'Returns { label, color } based on points balance (Safe/Warning/Final Warning/Suspension Review/Suspended).'),
    ]),
    ('lib/utils.ts', 'Utility function for merging Tailwind CSS class names.', [
        ('Exports', 'cn'),
        ('cn', 'Combines clsx and tailwind-merge to merge class names, handling conflicts by keeping the last Tailwind class.'),
        ('Dependencies', 'clsx, tailwind-merge'),
    ]),
    # Types
    ('types/index.ts', 'TypeScript type definitions for all Firestore data models and application entities.', [
        ('Interfaces', 'Driver, Licence, PoliceOfficer, OffenceCategory, OffenceRecord, Appeal, Suspension, AuditLog, Notification, SystemSetting, Vehicle, Payment, QRCodeData, AppUser'),
        ('Types', 'UserRole (\'admin\' | \'police\' | \'driver\' | \'authority\')'),
        ('Usage', 'Imported by services, hooks, contexts, and page components for type safety.'),
    ]),
    # Firebase
    ('firebase/config.ts', 'Firebase SDK initialization and service instance exports.', [
        ('Exports', 'app, auth, db, storage, functions'),
        ('Initialization', 'Initializes Firebase app with hardcoded config. Conditionally initializes only on client side (typeof window !== \'undefined\').'),
        ('Services', 'Auth (email/password), Firestore (NoSQL database), Storage (file uploads), Functions (serverless)'),
        ('Important Note', 'Firebase configuration values are hardcoded in this file rather than loaded from environment variables.'),
    ]),
    # Seed
    ('seed/seed.ts', 'Database seed script populating the entire Firebase project with demo data.', [
        ('Purpose', 'Creates 14 Firebase Auth users with associated Firestore documents for development and demo.'),
        ('Data Created', '14 auth users, 14 driver profiles with licences, 2 police officers, 14 offence categories, 20+ sample offences, 1 appeal, 10 system settings, 8 regions, 8 notifications.'),
        ('Execution', 'npm run seed (via tsx)'),
        ('Dependencies', 'firebase/app, firebase/firestore, firebase/auth'),
    ]),
    ('seed/add-offences.ts', 'Supplementary seed script adding matatu/nganya offence categories.', [
        ('Purpose', 'Adds 2 specific offence categories for public service vehicles (EXC-01 Excess Passengers, HNG-01 Hanging Outside).'),
        ('Execution', 'npm run seed:offences (via tsx)'),
        ('Safety', 'Checks for existing categories before adding to avoid duplicates.'),
        ('Dependencies', 'firebase/app, firebase/firestore, firebase/auth'),
    ]),
]

for fname, purpose, details in files_doc:
    doc.add_heading(f'4.{files_doc.index((fname, purpose, details))+1} {fname}', level=2)
    doc.add_paragraph(f'Purpose: {purpose}')
    for label, content in details:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        p.add_run(str(content))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. FUNCTION DOCUMENTATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Function Documentation', level=1)
doc.add_paragraph('This section documents every function found throughout the codebase.')

functions_doc = [
    # useFirestore.ts
    ('useCollection', 'hooks/useFirestore.ts', 'Real-time Firestore collection subscription hook.', [
        ('Parameters', 'collectionName: string, ...constraints: QueryConstraint[]'),
        ('Returns', '{ data: T[], loading: boolean, error: string | null }'),
        ('Logic', 'Creates a Firestore query (with optional constraints), subscribes via onSnapshot. On snapshot, maps docs to array with id. On error, sets error state. Cleans up listener on unmount or collectionName change.'),
        ('Error Handling', 'Catches Firestore errors and sets error state string.'),
        ('Side Effects', 'Establishes real-time listener on Firestore. Cleans up on unmount.'),
        ('Called by', 'All pages that display Firestore data (dashboards, offences, appeals, drivers, licence, analytics, notifications, settings).'),
        ('Example', 'const { data: offences } = useCollection(\'offences\');'),
    ]),
    ('useDocument', 'hooks/useFirestore.ts', 'Real-time Firestore document subscription hook.', [
        ('Parameters', 'collectionName: string, docId: string | undefined'),
        ('Returns', '{ data: T | null, loading: boolean }'),
        ('Logic', 'If docId is undefined, returns null immediately. Otherwise subscribes via onSnapshot.'),
        ('Called by', 'Not directly called by any page. Available for future use.'),
    ]),
    ('addDocument', 'hooks/useFirestore.ts', 'Adds a new document to a Firestore collection.', [
        ('Parameters', 'collectionName: string, data: DocumentData'),
        ('Returns', 'Promise<DocumentReference> from addDoc'),
        ('Logic', 'Adds createdAt and updatedAt timestamps to the data object, then calls addDoc.'),
        ('Error Handling', 'Throws Firestore errors. Callers must wrap in try/catch.'),
        ('Called by', 'NewOffencePage (handleSubmit), AppealsPage (handleSubmit), SettingsPage (handleSave), NotificationsPage (markAllRead via dynamic import)'),
    ]),
    ('updateDocument', 'hooks/useFirestore.ts', 'Updates an existing Firestore document.', [
        ('Parameters', 'collectionName: string, docId: string, data: Partial<DocumentData>'),
        ('Returns', 'Promise<void>'),
        ('Logic', 'Merges updatedAt timestamp into data, calls updateDoc.'),
        ('Called by', 'SettingsPage, NotificationsPage, AppealsPage (via dynamic imports)'),
    ]),
    ('deleteDocument', 'hooks/useFirestore.ts', 'Deletes a Firestore document.', [
        ('Parameters', 'collectionName: string, docId: string'),
        ('Returns', 'Promise<void>'),
        ('Logic', 'Calls deleteDoc on the specified document reference.'),
        ('Called by', 'Not called by any page. Available for future use.'),
    ]),
    ('getDocument', 'hooks/useFirestore.ts', 'Fetches a single Firestore document (one-time, not real-time).', [
        ('Parameters', 'collectionName: string, docId: string'),
        ('Returns', 'Promise<T | null>'),
        ('Logic', 'Calls getDoc. Returns object with id and data if exists, null otherwise.'),
        ('Called by', 'NewOffencePage (lookupDriver function)'),
    ]),
    # authService.ts
    ('onAuthChange', 'services/authService.ts', 'Subscribes to Firebase auth state changes.', [
        ('Parameters', 'cb: (user: User | null) => void'),
        ('Returns', 'Unsubscribe function'),
        ('Logic', 'Passes callback to onAuthStateChanged.'),
        ('Called by', 'AuthContext useEffect'),
    ]),
    ('login', 'services/authService.ts', 'Authenticates user with email and password.', [
        ('Parameters', 'email: string, password: string'),
        ('Returns', 'Promise<UserCredential>'),
        ('Logic', 'Calls signInWithEmailAndPassword.'),
        ('Error Handling', 'Throws Firebase auth errors (wrong password, user not found, etc.)'),
        ('Called by', 'AuthPage (handleLogin, quickLogin)'),
    ]),
    ('register', 'services/authService.ts', 'Creates a new user account with profile data.', [
        ('Parameters', 'email: string, password: string, role: UserRole, profile: Record<string, any>'),
        ('Returns', 'Promise<UserCredential>'),
        ('Logic', 'Calls createUserWithEmailAndPassword, then writes user document to Firestore.'),
        ('Called by', 'Not called by any page. Available for registration functionality.'),
    ]),
    ('logout', 'services/authService.ts', 'Signs out the current user.', [
        ('Parameters', 'None'),
        ('Returns', 'Promise<void>'),
        ('Logic', 'Calls signOut.'),
        ('Called by', 'AuthContext logout function, Sidebar logout button'),
    ]),
    ('getUserProfile', 'services/authService.ts', 'Fetches user profile from Firestore.', [
        ('Parameters', 'uid: string'),
        ('Returns', 'Promise<AppUser | null>'),
        ('Logic', 'Reads users/{uid} document. Returns data cast as AppUser if exists, null otherwise.'),
        ('Called by', 'AuthProvider useEffect'),
    ]),
    # driverService.ts
    ('searchDrivers', 'services/driverService.ts', 'Searches drivers by name prefix.', [
        ('Parameters', 'term: string'),
        ('Returns', 'Promise<Driver[]>'),
        ('Logic', 'Queries drivers collection with >= and <= \uf8ff pattern for prefix matching.'),
        ('Called by', 'Not called by any page. Available for search functionality.'),
    ]),
    ('getDriverByLicence', 'services/driverService.ts', 'Finds a driver by licence number.', [
        ('Parameters', 'licenceNumber: string'),
        ('Returns', 'Promise<Driver | null>'),
        ('Logic', 'Queries drivers collection where licenceNumber == value.'),
        ('Called by', 'Not called by any page. Available for licence verification.'),
    ]),
    # offenceService.ts
    ('createOffence', 'services/offenceService.ts', 'Creates a new offence record.', [
        ('Parameters', 'data: Omit<OffenceRecord, \'id\'>'),
        ('Returns', 'Promise<DocumentReference>'),
        ('Logic', 'Adds current ISO timestamp and calls addDoc.'),
        ('Called by', 'Not called by any page (the app uses useFirestore.addDocument directly in NewOffencePage).'),
    ]),
    ('getOffencesByDriver', 'services/offenceService.ts', 'Gets all offences for a specific driver.', [
        ('Parameters', 'driverId: string'),
        ('Returns', 'Promise<OffenceRecord[]>'),
        ('Logic', 'Queries offences collection filtered by driverId, ordered by timestamp descending.'),
        ('Called by', 'Not called by any page (app uses useCollection with filters).'),
    ]),
    ('getOffenceCategories', 'services/offenceService.ts', 'Gets all offence categories.', [
        ('Parameters', 'None'),
        ('Returns', 'Promise<OffenceCategory[]>'),
        ('Logic', 'Reads all documents from offenceCategories collection.'),
        ('Called by', 'Not called by any page (app uses useCollection).'),
    ]),
    # format.ts
    ('formatCurrency', 'lib/format.ts', 'Formats a number as Kenyan Shillings.', [
        ('Parameters', 'amount: number'),
        ('Returns', 'string (e.g., "KES 10,000")'),
        ('Logic', 'Uses Intl.NumberFormat with en-KE locale.'),
        ('Called by', 'AdminDashboard, AuthorityDashboard, DriverDashboard, OffencesPage, AnalyticsPage, ReportsPage'),
    ]),
    ('formatDate', 'lib/format.ts', 'Formats a date string or Date object.', [
        ('Parameters', 'date: string | Date'),
        ('Returns', 'string (e.g., "Jun 1, 2023")'),
        ('Logic', 'Uses Intl.DateTimeFormat with en-KE locale.'),
        ('Called by', 'LicencePage, ReportsPage'),
    ]),
    ('formatDateTime', 'lib/format.ts', 'Formats a date with time.', [
        ('Parameters', 'date: string | Date'),
        ('Returns', 'string (e.g., "Jun 1, 2023, 02:30 PM")'),
        ('Logic', 'Uses Intl.DateTimeFormat with en-KE locale including hour/minute.'),
        ('Called by', 'AdminDashboard, DriverDashboard, PoliceDashboard, OffencesPage, AppealsPage, NotificationsPage'),
    ]),
    ('getStatusColor', 'lib/format.ts', 'Returns Tailwind CSS classes for status badges.', [
        ('Parameters', 'status: string'),
        ('Returns', 'string (Tailwind classes)'),
        ('Called by', 'Not directly called (Badge component variant prop is used instead). Available for custom styling.'),
    ]),
    ('getRiskColor', 'lib/format.ts', 'Returns text color class based on risk score.', [
        ('Parameters', 'score: number'),
        ('Returns', 'string (text-green-600, text-yellow-600, or text-red-600)'),
        ('Called by', 'DriversPage'),
    ]),
    ('getRiskLabel', 'lib/format.ts', 'Returns human-readable risk level label.', [
        ('Parameters', 'score: number'),
        ('Returns', 'string ("Low", "Medium", "High", or "Unknown")'),
        ('Called by', 'DriversPage'),
    ]),
    ('getDemeritStatus', 'lib/format.ts', 'Returns demerit status label and color based on points.', [
        ('Parameters', 'points: number'),
        ('Returns', '{ label: string, color: string }'),
        ('Logic', 'Thresholds: >=15 Safe, >=10 Warning, >=5 Final Warning, >=1 Suspension Review, <1 Suspended.'),
        ('Called by', 'DriverDashboard'),
    ]),
    ('cn', 'lib/utils.ts', 'Merges Tailwind CSS class names with conflict resolution.', [
        ('Parameters', '...inputs: ClassValue[]'),
        ('Returns', 'string'),
        ('Logic', 'Uses clsx to combine classes, then tailwind-merge to resolve conflicts.'),
        ('Called by', 'Sidebar, TopBar, FirstTimeTutorial, all UI components (Button, Card, Badge, Input, Select, Table, Avatar)'),
    ]),
    # MapPicker
    ('init (MapPicker)', 'components/MapPicker.tsx', 'Initializes Leaflet map with marker and event handlers.', [
        ('Parameters', 'None (uses refs and props via closure)'),
        ('Returns', 'void'),
        ('Logic', 'Dynamically imports Leaflet and CSS. Creates map centered on Nairobi (-1.2921, 36.8219). Adds OpenStreetMap tiles. Creates draggable marker. Sets up dragend and click handlers. Calls updateValue with default coordinates.'),
        ('Error Handling', 'Guards against double initialization via initStarted ref. Uses destroyed flag for cleanup.'),
        ('Side Effects', 'Creates DOM elements in mapRef. Adds event listeners to map and marker.'),
        ('Cleanup', 'Removes map instance on component unmount.'),
    ]),
    ('updateValue (MapPicker)', 'components/MapPicker.tsx', 'Updates the GPS coordinate value and fetches reverse-geocoded address.', [
        ('Parameters', 'lat: number, lng: number'),
        ('Returns', 'Promise<void>'),
        ('Logic', 'Calls onChange prop with formatted lat/lng. Fetches address from Nominatim OpenStreetMap API. Silently catches errors.'),
        ('External API', 'Nominatim reverse geocoding API (openstreetmap.org).'),
        ('Error Handling', 'Silent catch — sets address to "Location captured" on failure.'),
    ]),
    ('locateMe (MapPicker)', 'components/MapPicker.tsx', 'Uses browser geolocation to center the map on user location.', [
        ('Parameters', 'None'),
        ('Returns', 'void'),
        ('Logic', 'Checks navigator.geolocation availability. Calls getCurrentPosition with high accuracy. Sets map view and marker position on success.'),
        ('Error Handling', 'Sets locating state to false on error.'),
        ('Timeouts', '10 seconds timeout for geolocation.'),
    ]),
    # AuthGuard etc.
    ('AuthGuard', 'components/AuthGuard.tsx', 'Protects routes by redirecting unauthenticated users.', []),
    # auth page
    ('handleLogin', 'app/auth/page.tsx', 'Handles form submission for authentication.', []),
    ('quickLogin', 'app/auth/page.tsx', 'One-click login with demo account credentials.', []),
    # licence page
    ('downloadPDF', 'app/licence/page.tsx', 'Downloads the digital licence card as a PDF file.', []),
    ('printLicence', 'app/licence/page.tsx', 'Opens a print-friendly HTML version of the licence in a new window.', []),
    # reports page
    ('generatePDF', 'app/reports/page.tsx', 'Generates a PDF with header, body table, and footer using jsPDF.', []),
    ('handleExportCSV', 'app/reports/page.tsx', 'Exports report data as a CSV file.', []),
    ('getReportData', 'app/reports/page.tsx', 'Returns report data structure based on selected report type.', []),
    # settings
    ('handleSave', 'app/settings/page.tsx', 'Saves all modified settings to Firestore.', []),
    # notifications
    ('markAllRead', 'app/notifications/page.tsx', 'Marks all notifications as read.', []),
    # appeals
    ('handleSubmit', 'app/appeals/page.tsx', 'Submits a new appeal to Firestore.', []),
]

for fname, loc, purpose, details in functions_doc:
    doc.add_heading(f'5.{functions_doc.index((fname, loc, purpose, details))+1} {fname}', level=2)
    p = doc.add_paragraph()
    run = p.add_run('File: ')
    run.bold = True
    p.add_run(loc)
    doc.add_paragraph(purpose)
    for label, content in details:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        p.add_run(str(content))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. API DOCUMENTATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. API Documentation', level=1)
doc.add_paragraph(
    'Note: RoadSafe360 has no backend API route handlers. All data operations occur directly between the '
    'React frontend and Firebase services via the Firebase Web SDK. This section documents the implicit '
    'API endpoints represented by Firebase service calls.'
)

doc.add_heading('6.1 Firebase Authentication Endpoints', level=2)
add_table(
    ['Operation', 'Firebase Method', 'Service File', 'Called From', 'Purpose'],
    [
        ['Sign In', 'signInWithEmailAndPassword', 'services/authService.ts (login)', 'app/auth/page.tsx', 'Authenticates user with email/password'],
        ['Sign Out', 'signOut', 'services/authService.ts (logout)', 'AuthContext, Sidebar', 'Ends user session'],
        ['Create User', 'createUserWithEmailAndPassword', 'services/authService.ts (register)', 'N/A (available)', 'Registers new user account'],
        ['Auth State', 'onAuthStateChanged', 'services/authService.ts (onAuthChange)', 'AuthContext', 'Subscribes to auth state changes'],
    ]
)

doc.add_heading('6.2 Firestore Collection Operations', level=2)
doc.add_paragraph('All Firestore CRUD operations are performed through hooks/useFirestore.ts.')

add_table(
    ['Collection', 'Read Method', 'Write Method', 'Security', 'Used By'],
    [
        ['users', 'onAuthStateChanged + getDoc', 'setDoc (register)', 'Authenticated users can read; admin/owner can write', 'AuthContext, authService'],
        ['drivers', 'useCollection + getDoc (lookup)', 'addDocument', 'Authenticated read; Admin/Police write', 'All dashboards, drivers page, new offence'],
        ['licences', 'useCollection', 'addDocument (seed only)', 'Authenticated read; Admin write', 'Admin/Authority dashboards, licence page'],
        ['offences', 'useCollection', 'addDocument', 'Authenticated read; Police/Admin create; Admin update/delete', 'Admin/Police/Driver dashboards, offences page'],
        ['offenceCategories', 'useCollection', 'addDocument (seed only)', 'Authenticated read; Admin write', 'New offence page'],
        ['appeals', 'useCollection', 'addDocument', 'Authenticated read; Driver/Admin create; Admin/Authority update', 'Appeals page, driver dashboard'],
        ['policeOfficers', 'useCollection', 'addDocument (seed only)', 'Authenticated read; Admin write', 'Admin dashboard, reports'],
        ['notifications', 'useCollection', 'addDocument (seed only)', 'Authenticated read; Admin write', 'Notifications page, driver dashboard'],
        ['settings', 'useCollection', 'updateDocument', 'Authenticated read; Admin write', 'Settings page'],
        ['regions', 'useCollection', 'addDocument (seed only)', 'Authenticated read; Admin/Authority write', 'Regions page'],
        ['vehicles', 'useCollection', 'N/A', 'Authenticated read; Admin/Police write', 'Licence page'],
    ]
)

doc.add_heading('6.3 Endpoint Dependency Map (Request Lifecycle)', level=2)
doc.add_paragraph('User Login Lifecycle:')
steps = [
    '1. User visits /auth (ClientLayout renders without DashboardShell)',
    '2. AuthProvider useEffect subscribes to onAuthStateChanged',
    '3. User enters credentials and submits form',
    '4. handleLogin calls login() from authService',
    '5. authService.login calls firebase/auth signInWithEmailAndPassword',
    '6. Firebase validates credentials against Firebase Authentication',
    '7. onAuthStateChanged fires with User object',
    '8. AuthProvider fetches profile via getUserProfile(uid) → Firestore read: users/{uid}',
    '9. AuthPage useEffect sees user + profile → router.replace to role-based dashboard',
    '10. DashboardShell renders with AuthGuard → Sidebar shows role-specific links',
    '11. Dashboard page mounts → useCollection hooks establish real-time Firestore listeners',
    '12. Data flows to Recharts charts, stat cards, and tables via React re-render',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph('Offence Issuance Lifecycle:')
steps2 = [
    '1. Police officer navigates to /offences/new',
    '2. Page loads with useCollection(\'offenceCategories\') — real-time category list',
    '3. Officer types Driver ID → debounced getDocument(\'drivers\', id) lookup',
    '4. Driver info displayed with verification badge',
    '5. Officer selects offence category from dropdown',
    '6. MapPicker lazy-loads Leaflet → click/drag sets GPS coordinates',
    '7. Officer enters notes (optional)',
    '8. Submit → addDocument(\'offences\', {...}) creates offence record in Firestore',
    '9. Firestore security rules validate request (isPolice() || isAdmin())',
    '10. Success toast shown → form resets → real-time listener updates offences list',
]
for s in steps2:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. COMPONENT DOCUMENTATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Component Documentation', level=1)

components_table = [
    ('AuthGuard', 'Route protection wrapper',
     '{ children: React.ReactNode }', 'None',
     'useAuth (user, profile, loading), useRouter',
     'useEffect: redirect to /auth if not authenticated',
     'SkeletonLoader', 'DashboardShell',
     'Tailwind via SkeletonLoader', 'Loading state change, auth state change'),
    ('ClientLayout', 'Root client-side layout',
     '{ children: React.ReactNode }', 'None',
     'usePathname', 'useEffect: unregister old service workers',
     'ThemeProvider, AuthProvider, DashboardShell, PWAInstallPrompt, Toaster', 'RootLayout in app/layout.tsx',
     'Tailwind classes', 'Pathname change'),
    ('DashboardShell', 'Dashboard wrapper with sidebar and topbar',
     '{ children: React.ReactNode }', 'mobileMenuOpen',
     'None (uses children prop)', 'None',
     'Sidebar, TopBar, FirstTimeTutorial, AuthGuard', 'ClientLayout',
     'Tailwind (flex layout, responsive)', 'None'),
    ('Sidebar', 'Navigation sidebar',
     '{ open: boolean, onClose: () => void }', 'None',
     'useAuth, usePathname', 'handleNav (closes on mobile nav)',
     'None', 'DashboardShell',
     'Tailwind (fixed positioning, transitions), CSS variables', 'Pathname change, role change'),
    ('TopBar', 'Header bar',
     '{ onMenuToggle, mobileMenuOpen }', 'None',
     'useTheme, useAuth, usePathname', 'Theme toggle click, menu toggle click',
     'None', 'DashboardShell',
     'Tailwind (fixed header, backdrop blur)', 'Theme change, pathname change'),
    ('MapPicker', 'Interactive map for GPS location',
     '{ value: string, onChange: (value: string) => void }', 'mapReady, address, locating',
     'useRef (map, marker, init guard)', 'Map click, marker drag, locateMe click',
     'Nominatim reverse geocode API', 'NewOffencePage',
     'Leaflet CSS, Tailwind', 'Coordinate change via drag/click/geolocation'),
    ('FirstTimeTutorial', 'Onboarding overlay',
     'None (uses useAuth)', 'step, visible',
     'useAuth', 'Step navigation, dismiss',
     'None', 'DashboardShell',
     'Tailwind (overlay, card, animations), CSS variables', 'Role change, step change'),
    ('PWAInstallPrompt', 'Install app banner',
     'None', 'deferredPrompt, showPrompt, isStandalone',
     'useState, useEffect', 'beforeinstallprompt event, install click',
     'None', 'ClientLayout',
     'Tailwind (fixed bottom, card)', 'Install event'),
    ('SkeletonLoader', 'Loading placeholder',
     'None', 'None',
     'None', 'None',
     'None', 'AuthGuard, AuthPage',
     'Tailwind (shimmer animation)', 'None'),
    ('ServiceWorkerRegister', 'SW registration',
     'None', 'None',
     'useEffect', 'Service worker registration',
     'None', 'Not imported',
     'None', 'None'),
]

add_table(
    ['Component', 'Purpose', 'Props', 'State', 'Hooks', 'Events', 'Child Components', 'Parent', 'Styling', 'Re-render Triggers'],
    [(c[0], c[1], c[2], c[3], c[4], c[5], c[6], c[7], c[8], c[9]) for c in components_table]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. DATABASE DOCUMENTATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Database Documentation', level=1)

doc.add_heading('8.1 Database Technology', level=2)
doc.add_paragraph(
    'RoadSafe360 uses Firebase Firestore as its primary database. Firestore is a NoSQL document database '
    'that stores data in collections of documents. Each document contains key-value pairs and supports '
    'subcollections. The database schema is also documented in a PostgreSQL-compatible SQL format in SCHEMA.md '
    'for reference, but the actual implementation uses Firestore.'
)

doc.add_heading('8.2 Collections and Documents', level=2)
add_table(
    ['Collection', 'Document ID', 'Key Fields', 'Description'],
    [
        ['users', 'Firebase Auth UID', 'uid, email, role, displayName, profileId', 'User profiles linked to Firebase Auth'],
        ['drivers', 'Auto-generated', 'nationalID, fullName, phoneNumber, email, photoURL, bloodGroup, emergencyContact, pointsBalance, status, riskScore', 'Driver records with demerit points and risk scores'],
        ['licences', 'Auto-generated', 'licenceNumber, licenceClass, issueDate, expiryDate, status, driverId, vehicleCategories, endorsements', 'Driver licence records linked to drivers'],
        ['vehicles', 'Auto-generated', 'registrationNumber, make, model, year, colour, vehicleClass, chassisNumber, engineNumber, driverId, status, insuranceExpiry, roadTaxExpiry', 'Vehicle registrations linked to drivers'],
        ['policeOfficers', 'Auto-generated', 'badgeNumber, name, email, region, assignedStation, userId', 'Police officer profiles'],
        ['offenceCategories', 'Auto-generated', 'code, name, description, fineAmount, demeritPoints, severity, applicableVehicleClasses, repeatOffenderMultiplier, courtRequired', 'Lookup table for offence types'],
        ['offences', 'Auto-generated', 'timestamp, gpsLocation, notes, pointsDeducted, fineAmount, status, driverId, offenceCategoryId, officerId, mediaUrls', 'Individual traffic offence records'],
        ['appeals', 'Auto-generated', 'submissionDate, reason, evidenceURL, status, offenceRecordId, driverId, resolution, resolvedAt', 'Driver appeals against offences'],
        ['suspensionHistory', 'Auto-generated', 'startDate, endDate, reason, durationMonths, driverId, suspensionCount', 'Record of licence suspensions'],
        ['notifications', 'Auto-generated', 'recipientID, type, title, message, isRead, timestamp', 'System notifications to users'],
        ['settings', 'Auto-generated', 'key, value, description', 'System configuration key-value pairs'],
        ['regions', 'Auto-generated', 'name, code, officerCount, driverCount, offenceCount, stations', 'Regional road safety statistics'],
        ['payments', 'Auto-generated', 'driverId, offenceRecordId, amount, status, paidAt', 'Fine payment records'],
        ['auditLogs', 'Auto-generated', 'actorID, actionType, timestamp, targetEntityID, details', 'Audit trail for compliance'],
        ['reports', 'Auto-generated', 'Varies by report type', 'Generated report storage'],
        ['demeritHistory', 'Auto-generated', 'Varies', 'Historical demerit point changes'],
    ]
)

doc.add_heading('8.3 Relationships', level=2)
doc.add_paragraph(
    'Firestore is a NoSQL database, so relationships are implicit rather than enforced via foreign keys:'
)
add_table(
    ['Relationship', 'Source Field', 'Target Collection', 'Cardinality'],
    [
        ['User → Driver', 'users.profileId', 'drivers', 'One-to-one (optional)'],
        ['User → Police Officer', 'users.profileId', 'policeOfficers', 'One-to-one (optional)'],
        ['Driver → Licence', 'licences.driverId', 'drivers', 'One-to-one'],
        ['Driver → Vehicle', 'vehicles.driverId', 'drivers', 'One-to-many'],
        ['Driver → Offence', 'offences.driverId', 'drivers', 'One-to-many'],
        ['Driver → Appeal', 'appeals.driverId', 'drivers', 'One-to-many'],
        ['Driver → Suspension', 'suspensionHistory.driverId', 'drivers', 'One-to-many'],
        ['Offence → OffenceCategory', 'offences.offenceCategoryId', 'offenceCategories', 'Many-to-one'],
        ['Offence → Appeal', 'appeals.offenceRecordId', 'offences', 'One-to-one (optional)'],
        ['Offence → Payment', 'payments.offenceRecordId', 'offences', 'One-to-one (optional)'],
        ['Police Officer → Offence', 'offences.officerId', 'policeOfficers', 'One-to-many (optional)'],
    ]
)

doc.add_heading('8.4 Security Rules (firestore.rules)', level=2)
doc.add_paragraph(
    'Firestore security rules enforce access control at the document level. Helper functions:'
)
add_table(
    ['Function', 'Logic'],
    [
        ['isAuthenticated()', 'request.auth != null'],
        ['isAdmin()', 'Authenticated AND user document role == \'admin\''],
        ['isPolice()', 'Authenticated AND user document role == \'police\''],
        ['isDriver()', 'Authenticated AND user document role == \'driver\''],
        ['isAuthority()', 'Authenticated AND user document role == \'authority\''],
        ['isOwner(userId)', 'request.auth.uid == userId'],
    ]
)

doc.add_paragraph('Collection-level access summary:')
add_table(
    ['Collection', 'Read', 'Write'],
    [
        ['users', 'Authenticated', 'Admin or owner'],
        ['drivers', 'Authenticated', 'Admin or Police'],
        ['licences', 'Authenticated', 'Admin only'],
        ['offences', 'Authenticated', 'Create: Police/Admin; Update/Delete: Admin'],
        ['appeals', 'Authenticated', 'Create: Driver/Admin; Update: Admin/Authority'],
        ['settings', 'Authenticated', 'Admin only'],
        ['regions', 'Authenticated', 'Admin or Authority'],
        ['auditLogs', 'Admin only', 'Admin only'],
        ['reports', 'Admin or Authority', 'Admin only'],
        ['vehicles', 'Authenticated', 'Admin or Police'],
        ['payments', 'Authenticated', 'Admin only'],
        ['notifications', 'Authenticated', 'Admin only'],
    ]
)

doc.add_heading('8.5 Data Flow Examples', level=2)
doc.add_paragraph('Offence creation data flow:')
code_block("""Frontend (NewOffencePage) → addDocument('offences', {
  driverId: 'abc123',
  offenceCategoryId: 'cat456',
  pointsDeducted: 6,
  fineAmount: 10000,
  gpsLocation: '-1.2921, 36.8219',
  notes: 'Speeding at 145 km/h in 80 zone',
  status: 'issued',
  timestamp: '2026-07-22T14:30:00.000Z'
}) → Firestore SDK → Security rules check isPolice() || isAdmin()
→ Firestore writes document → Real-time listener triggers re-render""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. BUSINESS LOGIC
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Business Logic', level=1)

biz_rules = [
    (
        'Demerit Point System',
        'Every driver starts with 20 demerit points. Offences deduct points based on severity. The system tracks the remaining balance and determines licence status.',
        'Implemented in demo data (seed.ts sets pointsBalance values) and displayed via getDemeritStatus() in lib/format.ts. The driver dashboard shows the points gauge.',
        'Simulates Kenya\'s Traffic Act demerit point system for driver accountability.'
    ),
    (
        'Demerit Status Thresholds',
        'Points determine status: >=15 Safe, >=10 Warning, >=5 Final Warning, >=1 Suspension Review, <1 Suspended.',
        'getDemeritStatus() in lib/format.ts:48-53. Used on driver dashboard and licence page.',
        'Provides graduated warnings before suspension, giving drivers opportunities to improve.'
    ),
    (
        'Role-Based Access Control',
        'Four roles (admin, police, driver, authority) have different permissions and dashboards.',
        'Sidebar component filters navLinks by role. Firestore security rules enforce write/read permissions per collection. AuthProvider exposes role to the component tree.',
        'Ensures appropriate data access: police can issue offences, drivers see their own data, admins manage everything, authorities view analytics.'
    ),
    (
        'Appeal Workflow',
        'Drivers can submit appeals against offences. Appeals go through statuses: submitted → under_review → approved/rejected.',
        'AppealsPage handles submission with addDocument. Status badges display current stage. Security rules allow Driver/Admin to create, Admin/Authority to update.',
        'Provides due process for disputed traffic violations.'
    ),
    (
        'Offence Issuance Verification',
        'Before issuing an offence, the officer must verify the driver exists in the system by looking up the driver ID.',
        'lookupDriver function in NewOffencePage uses getDocument to fetch driver data. Show verified driver info before enabling submit button.',
        'Prevents issuing offences to non-existent drivers and ensures accurate record-keeping.'
    ),
    (
        'Risk Scoring',
        'Drivers have a riskScore between 0 and 1 based on offence history. Scores >=0.6 are "High", >=0.3 "Medium", <0.3 "Low".',
        'riskScore is a seed data field. getRiskLabel() and getRiskColor() in lib/format.ts provide display values. Authority dashboard shows risk prediction overview.',
        'Identifies high-risk drivers for targeted enforcement and intervention.'
    ),
    (
        'Licence Status Propagation',
        'When a driver\'s points reach 0, their licence status becomes suspended. Suspension is displayed on the digital licence card with visual warnings.',
        'Licence page checks licence.status and driver.status for suspension state. Shows red card styling, SUSPENDED badge, and warning icons.',
        'Provides clear visual communication of licence status to drivers and enforcement officers.'
    ),
    (
        'Quick Login for Demo',
        'The login page provides one-click buttons for demo accounts representing each role.',
        'quickLogin function in auth page passes hardcoded demo credentials to login(). Titled buttons show user name and role.',
        'Enables rapid testing and demonstration of all role-specific features without memorizing credentials.'
    ),
    (
        'Sustainable Offence Categories',
        'Offence categories are stored in Firestore and dynamically loaded, allowing admins to add/modify categories without code changes.',
        'useCollection(\'offenceCategories\') in NewOffencePage loads categories into a select dropdown. seed.ts and add-offences.ts populate initial data.',
        'Allows the system to accommodate new traffic regulations and offence types without development effort.'
    ),
]

for i, (name, rule, impl, reason) in enumerate(biz_rules, 1):
    doc.add_heading(f'9.{i} {name}', level=2)
    doc.add_paragraph(f'Rule: {rule}')
    p = doc.add_paragraph()
    run = p.add_run('Implementation: ')
    run.bold = True
    p.add_run(impl)
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Rationale: ')
    run2.bold = True
    p2.add_run(reason)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. AUTHENTICATION AND SECURITY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Authentication and Security', level=1)

doc.add_heading('10.1 Login Flow', level=2)
doc.add_paragraph(
    '1. User navigates to /auth. 2. If already authenticated (user + profile in AuthContext), useEffect '
    'redirects to role-specific dashboard. 3. User enters email and password (or clicks quick-login). '
    '4. handleLogin calls login() → signInWithEmailAndPassword(). 5. On success, AuthProvider\'s '
    'onAuthStateChanged fires. 6. getUserProfile fetches Firestore users document. 7. useEffect redirects '
    'to role dashboard. 8. On error, toast displays error message.'
)

doc.add_heading('10.2 Token Generation', level=2)
doc.add_paragraph(
    'Token generation is handled entirely by Firebase Authentication. The Firebase SDK manages ID tokens '
    'automatically — refreshing them as needed. No custom JWT generation exists in the application.'
)

doc.add_heading('10.3 Session Handling', level=2)
doc.add_paragraph(
    'Firebase Auth manages sessions via ID tokens stored in browser storage (IndexedDB). The onAuthStateChanged '
    'listener in AuthContext detects session state changes. On page reload, the persisted session is restored '
    'automatically by the Firebase SDK. Logout calls signOut() which clears the session.'
)

doc.add_heading('10.4 Password Handling', level=2)
doc.add_paragraph(
    'Password handling is entirely delegated to Firebase Authentication. The application never stores or '
    'transmits passwords except through the Firebase Auth SDK. Password reset functionality is not implemented '
    'in the current codebase.'
)

doc.add_heading('10.5 Permissions and Guards', level=2)
add_table(
    ['Mechanism', 'Location', 'Description'],
    [
        ['AuthGuard Component', 'components/AuthGuard.tsx', 'Redirects unauthenticated users to /auth'],
        ['Role-Based Navigation', 'components/Sidebar.tsx', 'Shows only role-permitted navigation links'],
        ['Firestore Security Rules', 'firestore.rules', 'Document-level read/write permissions per collection'],
        ['AuthContext Role', 'contexts/AuthContext.tsx', 'Provides role to entire component tree via React context'],
        ['ClientLayout Routing', 'components/ClientLayout.tsx', 'Only wraps non-auth pages in DashboardShell (which includes AuthGuard)'],
    ]
)

doc.add_heading('10.6 Rate Limiting', level=2)
doc.add_paragraph(
    'Rate limiting is not implemented in the application code. Firebase Authentication provides built-in '
    'rate limiting for login attempts at the infrastructure level.'
)

doc.add_heading('10.7 Security Headers', level=2)
doc.add_paragraph(
    'Security headers are not explicitly configured in next.config.ts. Next.js provides some default security '
    'headers in production builds. No Content-Security-Policy, X-Frame-Options, or other headers are customized.'
)

doc.add_heading('10.8 Input Validation', level=2)
doc.add_paragraph(
    'Input validation is minimal: (1) Form inputs use the HTML \'required\' attribute for basic validation. '
    '(2) Driver ID lookup validates that the driver exists before submission. (3) The debounce on driver lookup '
    'prevents excessive Firestore reads. (4) No Zod or react-hook-form validation is actively used in page forms, '
    'though both are dependencies. (5) No server-side validation exists as there is no backend.'
)

doc.add_heading('10.9 Sanitization', level=2)
doc.add_paragraph(
    'No explicit sanitization is performed. The application relies on React\'s built-in XSS protection (JSX '
    'escapes values) and Firebase SDK\'s parameterized queries. User-supplied values (notes, appeal reasons) '
    'are stored in Firestore as-is.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. ENVIRONMENT VARIABLES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. Environment Variables', level=1)

doc.add_paragraph(
    'Environment variables are defined in .env.example. However, the Firebase configuration is hardcoded '
    'in firebase/config.ts with actual values, bypassing the environment variables. The seed scripts '
    'reference these environment variables as fallbacks:'
)

add_table(
    ['Variable', 'Purpose', 'Required', 'Default', 'Files Used', 'Effect'],
    [
        ['NEXT_PUBLIC_FIREBASE_API_KEY', 'Firebase API key for authentication', 'Yes', 'AIzaSyBHkqVST88k1Ojdx_96QWbnjky3-xnwBF8', 'firebase/config.ts, seed/seed.ts, seed/add-offences.ts', 'Firebase initialization fails if invalid'],
        ['NEXT_PUBLIC_FIREBASE_AUTH_DOMAIN', 'Firebase Auth domain', 'Yes', 'roadsafe360-95d87.firebaseapp.com', 'firebase/config.ts, seed/seed.ts', 'Auth operations fail if incorrect'],
        ['NEXT_PUBLIC_FIREBASE_PROJECT_ID', 'Firebase project identifier', 'Yes', 'roadsafe360-95d87', 'firebase/config.ts, seed/seed.ts', 'All Firebase service access fails if incorrect'],
        ['NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET', 'Firebase Storage bucket URL', 'Yes', 'roadsafe360-95d87.firebasestorage.app', 'firebase/config.ts, seed/seed.ts', 'File upload/download fails if incorrect'],
        ['NEXT_PUBLIC_FIREBASE_MESSAGING_SENDER_ID', 'Firebase Cloud Messaging sender', 'Yes', '64803316056', 'firebase/config.ts, seed/seed.ts', 'Cloud Messaging fails if needed'],
        ['NEXT_PUBLIC_FIREBASE_APP_ID', 'Firebase application identifier', 'Yes', '1:64803316056:web:5814a4f41f6b467123452d', 'firebase/config.ts, seed/seed.ts', 'Firebase SDK initialization fails if incorrect'],
    ]
)

doc.add_paragraph(
    'IMPORTANT: The Firebase config in firebase/config.ts is hardcoded with live credentials. Seed scripts '
    'use environment variables as fallbacks but also have hardcoded defaults. This is a security concern — '
    'the hardcoded values should be removed and only loaded from environment variables.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 12. CONFIGURATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12. Configuration', level=1)

config_files = [
    ('next.config.ts', 'Next.js framework configuration.', [
        ('images.remotePatterns', 'Whitelists Firebase Storage (firebasestorage.googleapis.com) and Google (lh3.googleusercontent.com) as allowed image sources for the Next.js Image component.'),
    ]),
    ('tsconfig.json', 'TypeScript compiler configuration.', [
        ('target', 'ES2017 — modern JavaScript output.'),
        ('module', 'esnext with bundler module resolution.'),
        ('jsx', 'react-jsx — the new JSX transform (no need to import React).'),
        ('strict', 'true — all strict type-checking options enabled.'),
        ('paths', '@/* maps to ./* — path alias for cleaner imports.'),
        ('include', 'All .ts, .tsx, .mts files plus .next/types.'),
    ]),
    ('eslint.config.mjs', 'ESLint configuration.', [
        ('Extends', 'eslint-config-next/core-web-vitals and typescript.'),
        ('Ignores', '.next/, out/, build/, next-env.d.ts.'),
    ]),
    ('postcss.config.mjs', 'PostCSS configuration for Tailwind CSS.', [
        ('Plugin', '@tailwindcss/postcss — the Tailwind CSS v4 PostCSS plugin.'),
    ]),
    ('public/manifest.json', 'PWA web app manifest.', [
        ('display', 'standalone — runs as a native-like app when installed.'),
        ('start_url', '/auth — opens login page on launch.'),
        ('theme_color', '#1e40af — matches the primary brand color.'),
        ('icons', '192x192 and 512x512 PNG icons for splash screens.'),
    ]),
    ('firestore.rules', 'Firestore security rules configuration.', [
        ('rules_version', '2 — latest Firestore rules syntax.'),
        ('Access Control', 'Role-based: isAuthenticated, isAdmin, isPolice, isDriver, isAuthority helper functions. Each collection has specific read/write rules.'),
    ]),
    ('storage.rules', 'Firebase Storage security rules.', [
        ('Default', 'Authenticated users can read all files. Write allowed for files < 10MB.'),
        ('Paths', 'Separate rules for /offences/, /driver-photos/, /appeal-evidence/ paths.'),
    ]),
    ('package.json', 'NPM package configuration.', [
        ('Scripts', 'dev (next dev), build (next build), start (next start), lint (eslint), seed (tsx seed/seed.ts), seed:offences (tsx seed/add-offences.ts).'),
        ('Dependencies', 'See Section 1.4 for detailed dependency breakdown.'),
    ]),
]

for fname, purpose, options in config_files:
    doc.add_heading(f'12.{config_files.index((fname, purpose, options))+1} {fname}', level=2)
    doc.add_paragraph(purpose)
    for key, val in options:
        p = doc.add_paragraph()
        run = p.add_run(f'{key}: ')
        run.bold = True
        p.add_run(val)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 13. ERROR DOCUMENTATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('13. Error Documentation', level=1)

doc.add_heading('13.1 Authentication Errors', level=2)
add_table(
    ['Error', 'Origin', 'Cause', 'Handling', 'User Experience', 'Recovery'],
    [
        ['auth/user-not-found', 'Firebase Auth (login)', 'Email not registered', 'Caught in handleLogin try/catch, toast displays error.message', 'Toast: "There is no user record..."', 'Check email or use quick-login button'],
        ['auth/wrong-password', 'Firebase Auth (login)', 'Incorrect password', 'Caught in handleLogin try/catch, toast displays error.message', 'Toast: "The password is invalid..."', 'Re-enter password or use quick-login button'],
        ['auth/invalid-credential', 'Firebase Auth (login)', 'Invalid email/password combination', 'Caught in handleLogin try/catch', 'Toast: "Invalid credentials"', 'Verify credentials or use quick-login'],
        ['auth/too-many-requests', 'Firebase Auth', 'Rate limited after multiple failures', 'Caught in try/catch', 'Toast: "Access to this account has been temporarily disabled"', 'Wait and try again later'],
        ['Auth profile null', 'AuthContext', 'User authenticated but no Firestore profile', 'setProfile(null), loading = false', 'AuthGuard redirects to /auth (no user/profile)', 'Ensure seed script created user document'],
    ]
)

doc.add_heading('13.2 Firestore Errors', level=2)
add_table(
    ['Error', 'Origin', 'Cause', 'Handling', 'User Experience', 'Recovery'],
    [
        ['Firestore permission denied', 'Firestore SDK', 'Security rules reject operation', 'Caught in try/catch, toast displays error.message', 'Toast: "Missing or insufficient permissions"', 'Verify user role has correct permissions for operation'],
        ['Firestore unavailable', 'Firestore SDK', 'Network failure or service outage', 'useCollection sets error state', 'Data remains stale or empty', 'Check network connection and Firebase console status'],
        ['Document not found', 'getDocument lookup', 'Invalid document ID', 'Driver lookup returns null', 'Red border on input, "Driver not found" message', 'Check the driver ID and try again'],
    ]
)

doc.add_heading('13.3 Runtime Errors', level=2)
add_table(
    ['Error', 'Origin', 'Cause', 'Handling', 'User Experience', 'Recovery'],
    [
        ['Cannot read properties of undefined', 'Various pages', 'Driver or licence data not yet loaded', 'Optional chaining (?.) or conditional rendering', 'SkeletonLoader shows while loading', 'Wait for data to load (Firestore real-time sync)'],
        ['pdf save failure', 'downloadPDF in licence page', 'html2canvas or jsPDF failure', 'Caught in try/catch, toast.error', 'Toast: "Failed to generate PDF"', 'Try again or use print option'],
        ['Pop-up blocked', 'printLicence', 'Browser blocked window.open', 'Check if popup is null, toast.error', 'Toast: "Please allow popups to print"', 'Allow popups for the site and try again'],
        ['Leaflet map init crash', 'MapPicker', 'Double initialization in StrictMode', 'initStarted ref prevents double init, destroyed flag cleanup', 'Map appears normally', 'Resolved — handled by the ref guard pattern'],
        ['Dynamic import failure', 'NewOffencePage (MapPicker)', 'Network error loading Leaflet', 'Suspense fallback shows shimmer', 'Shimmer placeholder instead of map', 'Check network, reload page'],
    ]
)

doc.add_heading('13.4 Development/Build Errors', level=2)
add_table(
    ['Error', 'Origin', 'Cause', 'Handling', 'Recovery'],
    [
        ['Module not found: Can\'t resolve \'@/...\'', 'Next.js build', 'Incorrect import path', 'Build fails', 'Verify path alias in tsconfig.json'],
        ['TypeScript type errors', 'Type-checking', 'Mismatched types', 'Compilation fails', 'Fix type definitions or casts'],
        ['Firebase not initialized', 'firebase/config.ts', 'Running on server side', 'Conditional init (typeof window check)', 'Ensure code runs only on client'],
        ['npm run seed fails', 'seed/seed.ts', 'Firebase Auth credentials wrong', 'Process exits with error', 'Verify Firebase project credentials'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 14. EXTERNAL SERVICES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('14. External Services', level=1)

add_table(
    ['Service', 'Purpose', 'Authentication', 'Configuration', 'Endpoints Used', 'Failure Handling', 'Retry'],
    [
        ['Firebase Authentication', 'User sign-in, session management', 'API key + project ID', 'Hardcoded in firebase/config.ts', 'signInWithEmailAndPassword, onAuthStateChanged, signOut', 'Errors caught in try/catch, displayed as toasts', 'Firebase SDK retries internally'],
        ['Firebase Firestore', 'NoSQL document database', 'Firebase Auth token', 'Hardcoded in firebase/config.ts', 'Collection/document reads/writes via SDK', 'Errors caught in try/catch or error state in hooks', 'None — Firestore SDK handles reconnection'],
        ['Firebase Storage', 'File storage for evidence/photos', 'Firebase Auth token', 'Hardcoded in firebase/config.ts', 'Upload/download via SDK', 'Not yet implemented (UI buttons exist)'],
        ['Firebase Functions', 'Serverless computing', 'Firebase Auth token', 'Hardcoded in firebase/config.ts', 'Not actively used', 'N/A'],
        ['OpenStreetMap (Nominatim)', 'Reverse geocoding for location addresses', 'None (free API)', 'URL in MapPicker updateValue', 'https://nominatim.openstreetmap.org/reverse', 'Silent catch — address set to "Location captured"', 'None'],
        ['OpenStreetMap (Tiles)', 'Map tile rendering', 'None (free tier)', 'URL in MapPicker init', 'https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png', 'Map loads without tiles if URL fails', 'None — tiles load on retry'],
        ['Leaflet CDN', 'Map marker assets', 'None', 'CDN URLs in MapPicker', 'https://unpkg.com/leaflet@1.9.4/dist/images/...', 'Markers may not render if CDN unavailable', 'None'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 15. DEPENDENCY GRAPH
# ═══════════════════════════════════════════════════════════════
doc.add_heading('15. Dependency Graph', level=1)

doc.add_heading('15.1 Module Dependencies', level=2)
doc.add_paragraph('Top-level dependency map:')

code_block("""app/ (pages)
  ├── components/ (AuthGuard, ClientLayout, DashboardShell, Sidebar, TopBar, etc.)
  ├── contexts/ (AuthContext, ThemeContext)
  ├── hooks/ (useFirestore)
  ├── services/ (authService, driverService, offenceService)
  ├── lib/ (utils.ts, format.ts)
  └── types/ (index.ts)

components/
  ├── contexts/ (AuthContext, ThemeContext)
  ├── hooks/ (useFirestore - via dynamic import in some components)
  ├── lib/ (utils.ts)
  └── types/ (via contexts)

contexts/
  ├── services/ (authService)
  ├── firebase/ (config.ts)
  └── types/ (index.ts)

hooks/
  └── firebase/ (config.ts)

services/
  ├── firebase/ (config.ts)
  └── types/ (index.ts)

seed/
  └── firebase/ (app, firestore, auth - standalone, not part of app)""")

doc.add_heading('15.2 Circular Dependencies', level=2)
doc.add_paragraph('No circular dependencies were detected. The dependency graph is acyclic and unidirectional.')

doc.add_heading('15.3 Shared Utilities', level=2)
doc.add_paragraph('The cn() function from lib/utils.ts is the most widely shared utility, used in all UI components and the Sidebar/FirstTimeTutorial. The format.ts functions are shared across all data-displaying pages.')

doc.add_heading('15.4 Core Services', level=2)
doc.add_paragraph('The core services that underpin all application functionality: (1) firebase/config.ts provides initialized Firebase instances. (2) hooks/useFirestore.ts provides all data access patterns. (3) contexts/AuthContext.tsx provides authentication state of the entire app.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 16. COMPLETE CALL GRAPH
# ═══════════════════════════════════════════════════════════════
doc.add_heading('16. Complete Call Graph', level=1)

doc.add_paragraph('Application Initialization:')
code_block("""RootLayout → ClientLayout → ThemeProvider → AuthProvider → [onAuthChange → getUserProfile]
         → DashboardShell (if not /auth) → AuthGuard → [redirect if no user]
         → Sidebar + TopBar + {page} + FirstTimeTutorial + PWAInstallPrompt + Toaster""")

doc.add_paragraph('User Login:')
code_block("""AuthPage: handleLogin/form submit → authService.login(email, password)
  → signInWithEmailAndPassword → Firebase Auth API
  ← onAuthStateChanged fires → AuthProvider: setUser
  → getUserProfile(uid) → getDoc(users/{uid}) → Firestore
  ← AuthPage useEffect → router.replace('/dashboard/{role}')
  → DashboardShell → Sidebar (role-based nav) + Dashboard page
  → useCollection hooks → onSnapshot → Firestore real-time
  → Recharts renders charts ← data state updates""")

doc.add_paragraph('Offence Issuance:')
code_block("""NewOffencePage: mount → useCollection('offenceCategories') → real-time list
  → handleDriverChange → debounce → getDocument('drivers', id)
  → MapPicker: lazy import → Leaflet init → marker/click handlers
  → handleSubmit → addDocument('offences', {...})
  → Firestore security rules → document written
  → toast.success → form reset""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 17. SEQUENCE DIAGRAMS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('17. Sequence Diagrams', level=1)

doc.add_heading('17.1 User Login Sequence', level=2)
code_block("""User          AuthPage       authService    Firebase Auth   AuthContext   Firestore
  │              │              │              │               │            │
  │─ enter creds │              │              │               │            │
  │─ submit ────>│              │              │               │            │
  │              │─ login() ───>│              │               │            │
  │              │              │─ signInWith  │               │            │
  │              │              │  EmailPw ───>│               │            │
  │              │              │              │─ validate ───>│            │
  │              │              │<─ token ─────│               │            │
  │              │              │<─ return ────│               │            │
  │              │              │              │─ onAuthChange │            │
  │              │              │              │─ setUser ────>│            │
  │              │              │              │              │─ getDoc ──>│
  │              │              │              │              │<─ profile ─│
  │              │<─ redirect ──│              │              │            │
  │─ dashboard ─>│              │              │              │            │
  │              │              │              │              │            │""")

doc.add_heading('17.2 Data Fetching (Real-time Subscription)', level=2)
code_block("""Component      useCollection    Firestore SDK    Firestore
  │              │                │               │
  │─ mount ─────>│                │               │
  │              │─ onSnapshot ──>│               │
  │              │                │─ listener ───>│
  │              │                │<─ snapshot ───│
  │              │─ setData ─────>│               │
  │<─ re-render ─│                │               │
  │              │                │               │
  │              │  (data changes on server)      │
  │              │                │<─ snapshot ───│
  │              │─ setData ─────>│               │
  │<─ re-render ─│                │               │
  │              │                │               │""")

doc.add_heading('17.3 Error Handling Flow', level=2)
code_block("""Component      Function        Firebase       Error Handler    Toast
  │              │               │              │               │
  │─ action ────>│               │              │               │
  │              │─ API call ───>│              │               │
  │              │               │─ error ─────>│               │
  │              │<─ throw ──────│              │               │
  │              │─ catch ─────────────────────>│               │
  │              │              │              │─ toast.error ─>│
  │<─ UI update ─│              │              │               │
  │              │              │              │               │""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 18. DEVELOPER GUIDE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('18. Developer Guide', level=1)

doc.add_heading('18.1 Required Software', level=2)
doc.add_paragraph('Node.js >= 18.x, npm >= 9.x, Git, A code editor (VS Code recommended)')

doc.add_heading('18.2 Installation', level=2)
code_block('''# Clone the repository
git clone <repository-url>
cd RoadSafe360

# Install dependencies
npm install

# Set up environment variables (optional - defaults exist in firebase/config.ts)
# Copy .env.example to .env.local and fill in Firebase values
cp .env.example .env.local''')

doc.add_heading('18.3 Running Locally', level=2)
code_block('''# Start development server
npm run dev

# Open http://localhost:3000 in browser
# Use any of the quick-login buttons on the auth page''')

doc.add_heading('18.4 Seeding the Database', level=2)
code_block('''# Seed the Firebase project with demo data
npm run seed

# Add matatu/nganya offence categories
npm run seed:offences''')

doc.add_heading('18.5 Running Tests', level=2)
doc.add_paragraph(
    'No test files exist in the current codebase. There are no test frameworks configured. '
    'To run linting only:'
)
code_block('npm run lint')

doc.add_heading('18.6 Building for Production', level=2)
code_block('''# Build the application
npm run build

# Start production server
npm start''')

doc.add_heading('18.7 Deployment Process', level=2)
doc.add_paragraph('Deployment is configured for Vercel (as specified in README.md):')
code_block('''# Deploy to Vercel (requires Vercel CLI or GitHub integration)
vercel --prod

# Ensure Firebase project is configured with matching credentials
# Update firestore.rules and storage.rules in Firebase Console''')

doc.add_heading('18.8 Common Development Workflow', level=2)
doc.add_paragraph(
    '1. Run npm run dev to start the development server. 2. Make changes to components in app/, components/, '
    'hooks/, or services/. 3. The application uses Hot Module Replacement (HMR) for instant updates. '
    '4. If you change Firestore data, the real-time listeners automatically update the UI. '
    '5. Run npm run lint before committing. 6. To reset demo data, re-run npm run seed.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 19. TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════
doc.add_heading('19. Troubleshooting', level=1)

add_table(
    ['Issue', 'Root Cause', 'Resolution', 'Debugging Tips', 'Logging Location'],
    [
        ['Cannot sign in with demo credentials', 'Firebase project credentials changed or seed not run', 'Run npm run seed to populate database', 'Check Firebase Console > Authentication for user list', 'Browser console > Firebase Auth errors'],
        ['Blank page on /auth', 'Firebase initialization failed (wrong config)', 'Check firebase/config.ts credentials match Firebase Console project settings', 'Open browser console for Firebase initialization errors', 'Browser console (F12 > Console tab)'],
        ['Offence list shows no data', 'Firestore collection is empty or security rules block access', 'Run npm run seed, verify Firestore security rules', 'Check Firebase Console > Firestore for collection existence', 'Browser console, Firestore usage tab'],
        ['Map does not appear on New Offence page', 'Leaflet dynamic import failed or CDN blocked', 'Check network tab for blocked requests to unpkg.com/openstreetmap.org', 'Network tab shows failed requests to Leaflet CDN / OpenStreetMap tiles', 'Browser console > Network tab'],
        ['PDF download fails silently', 'html2canvas or jsPDF error', 'Try using Print instead of PDF download', 'Wrap the downloadPDF in try/catch with console.log', 'Browser console error output'],
        ['Theme toggle not working', 'localStorage blocked or JavaScript disabled', 'Check browser privacy settings', 'Test localStorage.setItem in browser console', 'Browser console > localStorage tab'],
        ['npm run seed fails with "email-already-in-use"', 'User accounts already exist in Firebase project', 'Delete users from Firebase Console Auth tab, or the script handles by signing in', 'Check Firebase Console > Authentication', 'Terminal output shows which users failed'],
        ['Build fails with TypeScript errors', 'Type mismatches or missing type definitions', 'Run npx tsc --noEmit to check all errors', 'Check console for specific file and line number', 'Terminal output during build'],
        ['Charts not rendering correctly', 'Recharts version incompatibility or data format issues', 'Check that data arrays have expected structure', 'console.log the data prop being passed to Recharts', 'Browser console'],
        ['Notifications not showing as read after marking', 'Firestore document update failed due to permissions', 'Check user role has permission to update notifications collection', 'Check Firestore security rules for notifications collection', 'Browser console > Firestore errors'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 20. FUTURE DEVELOPMENT NOTES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('20. Future Development Notes', level=1)

doc.add_heading('20.1 Technical Debt', level=2)
items = [
    'Hardcoded Firebase credentials in firebase/config.ts — should use environment variables exclusively.',
    'Duplicate cn() function in utils/cn.ts (identical to lib/utils.ts). The file should be deleted.',
    'ServiceWorkerRegister component registers /sw.js but no service worker file exists in public/.',
    'createOffence function in offenceService.ts exists but the app uses useFirestore.addDocument directly (inconsistency).',
    'Driver search on PoliceDashboard is a visual-only input — actual search/filter is not implemented.',
    'Evidence and photo upload buttons on NewOffencePage exist in the UI but have no upload logic.',
    'No test files — the project has zero test coverage.',
    'Any type used extensively (over 40 occurrences) — type safety could be improved.',
    'No loading states on most data-fetching pages (the app relies on SkeletonLoader at the AuthGuard level).',
    'Interactive map placeholder on RegionsPage shows "Configure API key in .env to enable" but no API key variable is documented.',
]
for item in items:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_heading('20.2 Dead Code', level=2)
items2 = [
    'utils/cn.ts — duplicate of lib/utils.ts, not imported anywhere.',
    'ServiceWorkerRegister component — not imported or used in the app.',
    'services/offenceService.ts createOffence function — not called by any component.',
    'services/offenceService.ts getOffencesByDriver — not called by any component.',
    'services/offenceService.ts getOffenceCategories — not called by any component.',
    'services/driverService.ts searchDrivers — not called by any component.',
    'services/driverService.ts getDriverByLicence — not called by any component.',
    'services/authService.ts register — not called by any component.',
    'hooks/useFirestore.ts useDocument — not called by any component.',
    'hooks/useFirestore.ts deleteDocument — not called by any component.',
    'app/dashboard/layout.tsx — pass-through layout with no logic.',
]
for item in items2:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_heading('20.3 Unused Imports', level=2)
items3 = [
    'components/PWAInstallPrompt.tsx and components/ServiceWorkerRegister.tsx are imported in different parts but may have unused paths.',
    'react-hook-form and @hookform/resolvers are in package.json but barely used across the app.',
    'zod is in package.json but not used for validation anywhere.',
    'Multiple Radix UI packages in package.json that may not all be actively used.',
]
for item in items3:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_heading('20.4 Performance Considerations', level=2)
items4 = [
    'All pages use client-side rendering (\'use client\') — no SSR/SSG benefits for most pages.',
    'useCollection creates a Firestore listener for every mount — pages with multiple useCollection calls create multiple listeners.',
    'No pagination on any data table — offences, drivers, and appeals pages load all documents into memory.',
    'Recharts charts process all data client-side — large datasets could impact performance.',
    'No memoization (useMemo, useCallback) used except in MapPicker\'s lookupDriver.',
]
for item in items4:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_heading('20.5 Security Improvements', level=2)
items5 = [
    'Move Firebase credentials to environment variables. Current hardcoded values in firebase/config.ts are a security risk.',
    'Implement input validation (Zod or similar) on all forms.',
    'Add rate limiting for login attempts (Firebase Auth provides basic protection but application-level safeguards are absent).',
    'Implement Content Security Policy headers.',
    'Add audit logging for sensitive operations (offence issuance, appeals resolution, settings changes).',
    'Consider adding server-side API routes with middleware for additional security layers.',
]
for item in items5:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_heading('20.6 Maintainability Recommendations', level=2)
items6 = [
    'Adopt a consistent data access pattern — either use service modules or useFirestore hooks, not both.',
    'Add unit tests for utility functions (format.ts, utils.ts) and integration tests for data flows.',
    'Implement proper error boundaries for each major page section.',
    'Add TypeScript strict null checks (no implicit any).',
    'Create a centralized API layer instead of calling Firestore directly from page components.',
    'Add proper form state management with react-hook-form (already a dependency).',
    'Implement pagination and search on the server side (via Firestore queries with limits).',
    'Add loading skeletons for each page instead of the generic full-screen loader.',
]
for item in items6:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Appendix A: Demo Credentials', level=1)
add_table(
    ['Name', 'Email', 'Password', 'Role'],
    [
        ['Aisha Abubakar', 'aisha@roadsafe360.go.ke', 'Admin123!', 'Admin'],
        ['Khalid Salad', 'khalid@roadsafe360.go.ke', 'Auth123!', 'Authority'],
        ['Mohammed Karshe', 'mohammed@roadsafe360.go.ke', 'Officer123!', 'Police'],
        ['Adnan Ali', 'adnan@roadsafe360.go.ke', 'Officer123!', 'Police'],
        ['Zaahid Abdulmalik', 'zaahid.driver@roadsafe360.go.ke', 'Driver123!', 'Driver (Reckless)'],
        ['Aisha Abubakar (Driver)', 'aisha.driver@roadsafe360.go.ke', 'Driver123!', 'Driver (Off-duty)'],
        ['Khalid Salad (Driver)', 'khalid.driver@roadsafe360.go.ke', 'Driver123!', 'Driver (Off-duty)'],
        ['Mohammed Karshe (Driver)', 'mohammed.driver@roadsafe360.go.ke', 'Driver123!', 'Driver (Off-duty)'],
        ['Adnan Ali (Driver)', 'adnan.driver@roadsafe360.go.ke', 'Driver123!', 'Driver (Off-duty)'],
        ['Naila Amour', 'naila@roadsafe360.go.ke', 'Driver123!', 'Driver'],
        ['Rania Bahlewa', 'rania@roadsafe360.go.ke', 'Driver123!', 'Driver'],
        ['Zeitun Hussein', 'zeitun@roadsafe360.go.ke', 'Driver123!', 'Driver'],
        ['Reyhan Fuad', 'reyhan@roadsafe360.go.ke', 'Driver123!', 'Driver'],
        ['Turq Mahamud', 'turq@roadsafe360.go.ke', 'Driver123!', 'Driver'],
        ['Fenz Abdisalam', 'fenz@roadsafe360.go.ke', 'Driver123!', 'Driver'],
        ['UmulKheir Aden', 'umulkheir@roadsafe360.go.ke', 'Driver123!', 'Driver'],
        ['Faizaan Mohammed', 'faizaan@roadsafe360.go.ke', 'Driver123!', 'Driver'],
        ['Abdulhameed Saleh', 'abdulhameed@roadsafe360.go.ke', 'Driver123!', 'Driver'],
    ]
)

doc.add_heading('Appendix B: Offence Categories', level=1)
add_table(
    ['Code', 'Name', 'Fine (KES)', 'Points', 'Severity', 'Court Required'],
    [
        ['SPD-01', 'Speeding (Excess)', '10,000', '6', 'Serious', 'No'],
        ['SPD-02', 'Speeding (Minor)', '5,000', '3', 'Moderate', 'No'],
        ['DUI-01', 'Driving Under Influence', '100,000', '10', 'Serious', 'Yes'],
        ['RDL-01', 'Reckless Driving', '30,000', '8', 'Serious', 'Yes'],
        ['NLD-01', 'No Licence', '20,000', '4', 'Moderate', 'No'],
        ['NIN-01', 'No Insurance', '25,000', '4', 'Moderate', 'No'],
        ['TSG-01', 'Texting While Driving', '5,000', '3', 'Moderate', 'No'],
        ['SNL-01', 'No Seatbelt', '2,000', '1', 'Minor', 'No'],
        ['RTL-01', 'Running Red Light', '8,000', '4', 'Moderate', 'No'],
        ['OVL-01', 'Overloading', '15,000', '3', 'Moderate', 'No'],
        ['PRK-01', 'Illegal Parking', '1,500', '1', 'Minor', 'No'],
        ['HLM-01', 'Illegal Modification', '10,000', '2', 'Minor', 'No'],
        ['EXC-01', 'Excess Passengers (Matatu)', '5,000', '3', 'Moderate', 'No'],
        ['HNG-01', 'Hanging Outside (Matatu)', '3,000', '3', 'Moderate', 'No'],
    ]
)

doc.add_heading('Appendix C: System Settings', level=1)
add_table(
    ['Key', 'Default Value', 'Description'],
    [
        ['demerit_safe_threshold', '15', 'Minimum points for Safe status'],
        ['demerit_warning_threshold', '10', 'Minimum points for Warning status'],
        ['demerit_final_warning_threshold', '5', 'Minimum points for Final Warning status'],
        ['demerit_suspension_review_threshold', '1', 'Minimum points for Suspension Review status'],
        ['demerit_max_points', '20', 'Starting demerit points'],
        ['suspension_first_duration', '3', 'First suspension duration in months'],
        ['suspension_second_duration', '6', 'Second suspension duration in months'],
        ['suspension_third_action', 'revoked', 'Action after third suspension'],
        ['appeal_review_days', '14', 'Days to review an appeal'],
        ['licence_renewal_reminder_days', '30', 'Days before expiry to send reminder'],
    ]
)

doc.add_heading('Appendix D: File Inventory', level=1)
add_table(
    ['Category', 'Count', 'Details'],
    [
        ['Source files (TS/TSX)', '48', 'Application code, components, hooks, services, types, contexts'],
        ['Configuration files', '10', 'JSON, config.ts, rules files, tsconfig, eslint, postcss, gitignore'],
        ['Environment files', '1', '.env.example'],
        ['Documentation', '4', 'README.md, SCHEMA.md, CLAUDE.md, AGENTS.md'],
        ['Seed scripts', '2', 'seed/seed.ts, seed/add-offences.ts'],
        ['Static assets', '14', 'Images, SVGs, favicons, manifest.json'],
        ['CSS files', '1', 'app/globals.css'],
        ['Test files', '0', 'No tests configured'],
    ]
)

# ── SAVE ──
output_path = 'C:\\Users\\karei\\RoadSafe360\\documentation\\RoadSafe360_Technical_Documentation.docx'
doc.save(output_path)
print(f'Documentation saved to: {output_path}')
